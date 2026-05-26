from __future__ import annotations

import csv
import html
import json
import re
from datetime import date, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\sheri\Documents\GitHub\kangoo-theme")
KEYWORD_ROOT = Path(r"C:\Users\sheri\Documents\GitHub\kangoo-keywords-research")
OUT_DIR = ROOT / "guides" / "blog-guides"

SITE = "https://kangoopouches.co.uk"

LINKS = {
    "home": f"{SITE}/",
    "all": f"{SITE}/product-category/nicotine-pouches/",
    "trial": f"{SITE}/product-category/99p-pouches/",
    "zyn": f"{SITE}/product-category/zyn/",
    "velo": f"{SITE}/product-category/velo/",
    "pablo": f"{SITE}/product-category/pablo/",
    "killa": f"{SITE}/product-category/killa/",
    "mint": f"{SITE}/mint-nicotine-pouches/",
    "berry": f"{SITE}/berry-nicotine-pouches/",
    "strong": f"{SITE}/strong-strength-nicotine-pouches/",
    "extra": f"{SITE}/extra-strong-strength-nicotine-pouches/",
    "finder": f"{SITE}/pouch-finder/",
    "compare": f"{SITE}/compare-pouches/",
    "strength_ladder": f"{SITE}/strength-ladder/",
    "flavour_explorer": f"{SITE}/flavour-explorer/",
    "velo_bright": f"{SITE}/product/velo/bright-spearmint-6mg/",
    "velo_freezing": f"{SITE}/product/velo/freezing-peppermint-11mg/",
    "velo_ruby": f"{SITE}/product/velo/ruby-berry-10mg/",
    "zyn_cool": f"{SITE}/product/zyn/cool-mint-mini-3mg/",
    "zyn_cherry": f"{SITE}/product/zyn/black-cherry-mini-3mg/",
    "zyn_citrus": f"{SITE}/product/zyn/citrus-mini-6mg/",
    "pablo_ice": f"{SITE}/product/nicotine-pouches/pablo-ice-cold-nicotine-pouches-24mg/",
    "pablo_grape": f"{SITE}/product/nicotine-pouches/pablo-grape-ice-nicotine-pouches-30mg/",
}

SOURCES = {
    "nhs": "https://www.nhs.uk/better-health/quit-smoking/ready-to-quit-smoking/quit-with-nicotine-replacement-therapies-nrt/",
    "gov_snus": "https://www.gov.uk/government/publications/excise-notice-476-tobacco-products-duty/excise-notice-476-tobacco-products-duty/",
    "gov_ads": "https://www.gov.uk/marketing-advertising-law",
    "asa_compare": "https://www.asa.org.uk/advice-online/comparisons-verifiability.html",
}

NICOTINE_NOTE = (
    "Nicotine is addictive. This article is for adults who already use nicotine products. "
    "Nicotine pouches are not licensed stop-smoking products in the UK and this content is not medical advice."
)

PUBLIC_TEXT_REPLACEMENTS = {
    "When writing public recommendations, avoid saying a product is the cheapest unless the claim is backed by current evidence. It is better to say 'from 99p trial pouches' or 'multi-buy options available on selected products'.": "For value comparisons, use verifiable wording such as 'from 99p trial pouches' or 'multi-buy options available on selected products' rather than relying on broad cheapest claims.",
    "Use this guide to move from broad searches like 'ZYN pouches' or 'ZYN UK' into practical choices by strength and flavour.": "Broad searches like 'ZYN pouches' or 'ZYN UK' become more useful when you compare specific strengths and flavours.",
    "Do not present high strength as automatically better. Explain who it suits and link to the strength ladder for comparison.": "High strength is not automatically better. It is more relevant for experienced adult nicotine users who already know they prefer a stronger pouch.",
    "The content should be careful not to encourage users to jump straight to the highest strength. Strength should match experience level.": "It is better to compare experience level before moving straight to the highest strength. Strength should match the user, not just the product headline.",
    "This is why Kangoo content should use nicotine pouches as the main product wording and use snus only when explaining search language or legal differences.": "That is why this article uses nicotine pouches as the main product wording and uses snus only when explaining search language or legal differences.",
    "Do not describe nicotine pouches as medical products or licensed stop-smoking products. NHS guidance says nicotine pouches are not licensed stop-smoking products in the UK.": "Nicotine pouches are not medical products or licensed stop-smoking products. NHS guidance says nicotine pouches are not licensed stop-smoking products in the UK.",
    "Kangoo content should compare format and shopping factors, then send users to product pages only for tobacco-free pouch browsing.": "This comparison focuses on format and shopping factors, then points adult pouch shoppers toward tobacco-free pouch browsing.",
    "Send shoppers to the ZYN category first, then use product links for reliable picks such as Cool Mint, Black Cherry and Citrus when live stock supports it.": "Start with the ZYN category first, then use product links for picks such as Cool Mint, Black Cherry and Citrus when live stock supports it.",
    "Kangoo should route this search to the full category, ZYN, VELO and comparison tools rather than making unsupported competitor claims.": "The clearest next step is to compare the full category, ZYN, VELO and comparison tools rather than relying on unsupported competitor claims.",
    "Route shoppers to the PABLO brand page first, then to strong and extra-strong category pages if strength is their main filter.": "Start with the PABLO brand page first, then use strong and extra-strong category pages if strength is the main filter.",
    "This guide should explain terminology and product differences without suggesting nicotine improves sport performance.": "The key is terminology and product differences, without suggesting nicotine improves sport performance.",
    "Use this guide to link to how-to content and lower-strength comparison routes rather than pushing strong products.": "How-to content and lower-strength comparison routes are more useful here than jumping straight to strong products.",
    "Use this guide to support how-to searches, then route shoppers toward strength and flavour pages.": "Start with how-to guidance, then compare strength and flavour pages when you are ready to choose.",
    "This search intent is practical, so the page should answer quickly before explaining placement, timing and disposal.": "This question is practical, so the answer is direct: pouches are placed under the lip, removed after use and disposed of responsibly.",
    "Send shoppers into the full category and pouch finder to compare format and strength together.": "Use the full category and pouch finder to compare format and strength together.",
    "Route users to the full category, 99p pouch range and finder depending on their intent.": "Use the full category, 99p pouch range and finder depending on what you are trying to compare.",
    "Use this guide as an educational bridge to the nicotine pouch category, not as medical advice.": "This is an educational comparison that can lead into the nicotine pouch category, but it is not medical advice.",
    "Route users to tobacco-free nicotine pouch education, then category and finder pages.": "Use tobacco-free nicotine pouch education first, then category and finder pages when you are ready to compare products.",
    "This guide should explain when to use finder, compare, strength ladder and flavour explorer tools.": "The finder, compare page, strength ladder and flavour explorer each help at a different stage of choosing.",
    "Use tool pages as the main internal links, then direct confident shoppers to the full product category.": "Use the finder, flavour explorer, strength ladder and compare page first, then move into the full product category when the choice is clearer.",
    "The page should connect discreet delivery copy with real buying decisions: strength, flavour, price and stock.": "Discreet delivery matters most when it is connected to real buying decisions: strength, flavour, price and stock.",
    "Use commercial category links early because this is close to purchase intent.": "Start with commercial category pages because this search is close to purchase intent.",
    "Kangoo product and delivery copy should state current packaging terms clearly on site.": "Check the current delivery page or checkout copy for the latest packaging terms.",
    "Use this guide to support tobacco-free category messaging and link to brand pages.": "Use the tobacco-free category and brand pages to compare current products.",
    "This is why a UK ecommerce page should be careful with wording. Use 'nicotine pouches' for tobacco-free products, and explain that shoppers may use the word snus informally.": "That is why UK shoppers should look for the wording 'nicotine pouches' when they want tobacco-free products, even if they use the word snus informally.",
    "For Kangoo content, use KILLA nicotine pouches as the main wording, then explain that some shoppers may search for KILLA snus even though tobacco-free pouches are a different product type.": "KILLA nicotine pouches is the clearest product wording, even though some shoppers may search for KILLA snus when they mean tobacco-free pouches.",
    "The KILLA category page should be the main destination for this post. Add supporting links to strong, extra strong and compare pouches pages.": "The KILLA category page is the best starting point, with strong, extra strong and compare pouches pages useful for narrowing the choice.",
    "If you are writing for UK shoppers, use 'snus' only to explain meaning, search language or legal differences. Use 'nicotine pouches' for Kangoo products and product category pages.": "For UK shoppers, 'snus' is best understood as search language or a legal comparison term. For Kangoo products and category pages, 'nicotine pouches' is the clearer wording.",
}


def public_text(text: str) -> str:
    for old, new in PUBLIC_TEXT_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text


def load_keyword_rows() -> dict[str, tuple[int, float, str]]:
    data: dict[str, tuple[int, float, str]] = {}
    for path in KEYWORD_ROOT.glob("*.csv"):
        if path.name == "kp-products.csv":
            continue
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.DictReader(handle)
            for row in reader:
                keyword = (row.get("Keyword") or "").strip().lower()
                if not keyword:
                    continue
                try:
                    volume = int(float(row.get("Volume") or 0))
                except ValueError:
                    volume = 0
                try:
                    difficulty = float(row.get("Keyword Difficulty") or 0)
                except ValueError:
                    difficulty = 0.0
                existing = data.get(keyword)
                if existing is None or volume > existing[0]:
                    data[keyword] = (volume, difficulty, path.name)
    return data


KW_DATA = load_keyword_rows()


def keyword_note(keyphrase: str) -> str:
    lookup = KW_DATA.get(keyphrase.lower())
    if not lookup:
        return "No exact Semrush row found in the trimmed exports; use as a strategic long-tail target."
    volume, difficulty, source = lookup
    return f"Semrush export: volume {volume:,}, KD {difficulty:g}, source {source}."


def slugify(text: str) -> str:
    text = text.lower().replace("&", "and")
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text


GUIDES = [
    {
        "title": "Cheap Nicotine Pouches UK: 99p Trials, Prices and Multi-Buy Tips",
        "focus": "cheap nicotine pouches uk",
        "meta_title": "Cheap Nicotine Pouches UK From 99p | Kangoo",
        "meta_description": "Compare cheap nicotine pouches in the UK, including 99p trial pouches, regular cans, multi-buy packs and smart ways to choose value without guessing.",
        "secondary": ["99p nicotine pouches", "nicotine pouches free sample", "best nicotine pouches", "nicotine pouches uk", "cheap snus uk"],
        "intent": "Commercial investigation",
        "links": [
            ("99p nicotine pouches", LINKS["trial"]),
            ("shop all nicotine pouches", LINKS["all"]),
            ("compare nicotine pouches", LINKS["compare"]),
            ("Kangoo Pouch Finder", LINKS["finder"]),
            ("multi-buy nicotine pouch range", LINKS["all"]),
        ],
        "summary": "Lead with Kangoo's 99p trial angle, then explain how to compare price, stock, strength and pack value without making absolute lowest-price claims.",
        "sections": [
            ("Quick answer", [
                "The cheapest way to try nicotine pouches at Kangoo is to start with the 99p pouch range when trial stock is available. These trial pouches are designed for sampling flavours and strengths before committing to a regular can or multi-buy pack.",
                "For repeat orders, the real value usually comes from comparing the full price, the pack size, and the unit price. A pouch that looks cheaper as a single can is not always the best repeat-buy option if another product has a stronger multi-buy tier.",
            ]),
            ("How to compare cheap nicotine pouches properly", [
                "Price is only one part of the decision. Adult nicotine users should compare the can price, strength, flavour family, pouch count, stock status and delivery threshold before choosing.",
                "On Kangoo, the practical route is simple: check the 99p pouch page for trial stock, use the nicotine pouch category for the full range, then compare pack pricing on product pages where multi-buy options are available.",
            ]),
            ("99p trials vs regular cans", [
                "A 99p trial pouch is best treated as a low-cost way to test a flavour or strength. It is useful when you are unsure whether mint, berry, citrus or an extra-strong pouch is right for your regular rotation.",
                "Regular cans are better for repeat buying because they are the products most likely to support pack pricing, broader stock levels and stronger brand choice across ZYN, VELO, KILLA and PABLO.",
            ]),
            ("Where Kangoo can win on value", [
                "Kangoo's strongest value story is not only the headline trial price. It is the combination of 99p trials, regular products from the main nicotine pouch range, free delivery over the threshold and pack pricing on selected products.",
                "This is a cleaner claim than saying Kangoo is always the cheapest. Retailer pricing changes, and any public price comparison should be dated, representative and easy for shoppers to verify.",
            ]),
        ],
        "faq": [
            ("Are 99p nicotine pouches available all the time?", "No. They depend on live stock and trial availability. Treat the 99p range as a rotating sample section."),
            ("What should I buy after a 99p trial?", "Use the same flavour or strength as a guide, then compare regular cans and pack-priced products in the full nicotine pouches category."),
            ("Can I compare Kangoo with supermarkets?", "Yes, but keep comparisons generic unless you have dated, verifiable evidence for a named retailer claim."),
        ],
    },
    {
        "title": "99p Nicotine Pouches UK: How Kangoo Trial Pouches Work",
        "focus": "99p nicotine pouches",
        "meta_title": "99p Nicotine Pouches UK | Trial Pouches at Kangoo",
        "meta_description": "Learn how Kangoo 99p nicotine pouch trials work, what to check before buying and how to move from a sample pouch to a regular order.",
        "secondary": ["99p pouches", "nicotine pouches free sample", "cheap nicotine pouches", "snus free sample", "trial nicotine pouches"],
        "intent": "Transactional and commercial",
        "links": [
            ("99p pouch trials", LINKS["trial"]),
            ("full nicotine pouch range", LINKS["all"]),
            ("Kangoo Pouch Finder", LINKS["finder"]),
            ("compare pouches", LINKS["compare"]),
            ("ZYN 99p style trial picks", LINKS["trial"]),
        ],
        "summary": "Explain the one-per-order trial model, how to choose a first trial, and where to go next.",
        "sections": [
            ("Quick answer", [
                "Kangoo's 99p nicotine pouch range is a trial section for adult nicotine users who want to test selected pouch flavours and strengths at a low entry price.",
                "Trial pouches are not a permanent guarantee. They change with stock, and availability can move quickly when a popular flavour or strength is added.",
            ]),
            ("How 99p pouch trials work", [
                "The idea is simple: choose a trial pouch, add it to your basket while stock is available, and use that first order to learn what style of pouch you prefer.",
                "A shopper might test a mint pouch for freshness, a berry pouch for sweetness, or a stronger pouch only if they already know they prefer a higher-intensity product.",
            ]),
            ("How to choose your first 99p pouch", [
                "Start with flavour if you are new to the category. Mint, spearmint and peppermint are common everyday choices, while berry and citrus are useful if you want something fruit-led.",
                "Then check strength. Lower strengths are usually better for a gentler first test, while strong and extra-strong options should be kept for experienced adult nicotine users.",
            ]),
            ("What to do after trying a 99p pouch", [
                "If you liked the flavour and strength, move into the full nicotine pouch range and compare regular cans from the same brand or flavour family.",
                "If you liked the flavour but not the strength, use the strength ladder or pouch finder to find a better match before buying a multi-pack.",
            ]),
        ],
        "faq": [
            ("Are 99p nicotine pouches the same as free samples?", "No. They are paid trial products at a low price, not free products."),
            ("Can I build a full order from 99p pouches?", "The range is intended for trials, so use it to sample first and then compare regular products for repeat orders."),
            ("Do 99p pouches include ZYN or VELO?", "The trial mix can include popular brands when stock is available, but the live category page is the source of truth."),
        ],
    },
    {
        "title": "Best Nicotine Pouches UK: Brands, Strengths and Flavours Compared",
        "focus": "best nicotine pouches uk",
        "meta_title": "Best Nicotine Pouches UK | Compare Brands and Strengths",
        "meta_description": "Compare the best nicotine pouch options in the UK by brand, strength, flavour, price and use case, with links to ZYN, VELO, PABLO and KILLA ranges.",
        "secondary": ["best nicotine pouches", "nicotine pouches uk", "compare nicotine pouches", "zyn nicotine pouches", "velo nicotine pouches"],
        "intent": "Commercial investigation",
        "links": [
            ("compare nicotine pouches", LINKS["compare"]),
            ("shop nicotine pouches", LINKS["all"]),
            ("ZYN nicotine pouches", LINKS["zyn"]),
            ("VELO nicotine pouches", LINKS["velo"]),
            ("PABLO nicotine pouches", LINKS["pablo"]),
            ("KILLA nicotine pouches", LINKS["killa"]),
        ],
        "summary": "Define 'best' by shopper need rather than a single winner: best for mint, fruit, strong, value and first trial.",
        "sections": [
            ("Quick answer", [
                "The best nicotine pouch is the one that matches your preferred strength, flavour, price point and brand style. A single overall winner is less useful than a clear comparison by use case.",
                "For Kangoo shoppers, ZYN and VELO are strong everyday comparison brands, PABLO is more relevant for experienced users looking at higher strengths, and KILLA is useful for bold flavour and strong pouch shoppers.",
            ]),
            ("Best for first-time comparison", [
                "Start with a familiar flavour family such as mint or berry, then choose a lower or balanced strength before stepping into stronger products.",
                "The Kangoo Pouch Finder is the easiest internal route because it asks about experience, strength preference and flavour direction before pointing to product options.",
            ]),
            ("Best for brand-led shoppers", [
                "ZYN is a useful starting point for shoppers comparing mini formats, mint profiles and fruit flavours. VELO is strong for broad flavour coverage, especially peppermint, spearmint, berry and citrus-led options.",
                "PABLO and KILLA should be positioned more carefully because many products in those ranges sit toward the stronger end of the market.",
            ]),
            ("Best for value", [
                "Value depends on more than the can price. Check 99p trial pouches for sampling, then review regular cans and pack pricing where available.",
                "When writing public recommendations, avoid saying a product is the cheapest unless the claim is backed by current evidence. It is better to say 'from 99p trial pouches' or 'multi-buy options available on selected products'.",
            ]),
        ],
        "faq": [
            ("What are the best nicotine pouches for beginners?", "Lower or balanced strengths in familiar flavours such as mint, spearmint or berry are usually easier to compare first."),
            ("Which brand should I compare first?", "ZYN and VELO are good starting points because they include popular flavours and multiple strengths."),
            ("Are strong pouches better?", "Not automatically. Stronger pouches suit experienced users, but the best choice depends on tolerance and preference."),
        ],
    },
    {
        "title": "Nicotine Pouch Strength Guide: 3mg to Extra Strong Explained",
        "focus": "nicotine pouch strength guide",
        "meta_title": "Nicotine Pouch Strength Guide | 3mg to Extra Strong",
        "meta_description": "Understand nicotine pouch strengths from light and balanced to strong and extra strong, with tips for comparing mg, pouch feel and Kangoo products.",
        "secondary": ["nicotine pouch strength", "3mg nicotine pouches", "strong nicotine pouches", "extra strong nicotine pouches", "zyn strengths"],
        "intent": "Informational and commercial",
        "links": [
            ("strength ladder", LINKS["strength_ladder"]),
            ("strong nicotine pouches", LINKS["strong"]),
            ("extra strong nicotine pouches", LINKS["extra"]),
            ("pouch finder", LINKS["finder"]),
            ("shop nicotine pouches", LINKS["all"]),
        ],
        "summary": "Explain strength bands and how shoppers should compare mg, format and experience level.",
        "sections": [
            ("Quick answer", [
                "Nicotine pouch strength is usually shown in milligrams, but the number is only part of the experience. Pouch size, moisture, flavour cooling, format and personal tolerance can all affect how strong a pouch feels.",
                "A useful shopping framework is light, balanced, strong and extra strong. Kangoo's strength ladder turns those bands into a simpler buying route.",
            ]),
            ("Light and lower-strength pouches", [
                "Lower-strength pouches, such as 1.5mg, 3mg or similar mini formats, are usually easier to compare if you want a gentler pouch feel.",
                "They are also useful when flavour matters more than intensity, especially for mint, berry or citrus shoppers who want an everyday style.",
            ]),
            ("Balanced and strong pouches", [
                "Balanced products sit in the middle: enough presence for regular adult nicotine users, without moving straight into extra-strong territory.",
                "Strong pouches are better suited to shoppers who already know their tolerance and want a more noticeable pouch.",
            ]),
            ("Extra strong pouches", [
                "Extra strong pouches should be treated as experienced-user products. They are not automatically better, and they are not a shortcut to finding the right pouch.",
                "If you are unsure, compare strong and extra-strong pages first, then use the pouch finder to narrow the choice by flavour and brand.",
            ]),
        ],
        "faq": [
            ("Is 3mg a low strength?", "In most pouch ranges, 3mg is treated as a lower-strength option."),
            ("What does extra strong mean?", "It usually means a higher nicotine content or stronger pouch feel, but the exact experience varies by product."),
            ("Should I choose by mg only?", "No. Compare mg alongside flavour, pouch size, moisture, brand and your own experience level."),
        ],
    },
    {
        "title": "Strongest Nicotine Pouches UK: Strong and Extra Strong Options Explained",
        "focus": "strongest nicotine pouches",
        "meta_title": "Strongest Nicotine Pouches UK | Strong and Extra Strong",
        "meta_description": "Compare strong and extra strong nicotine pouches in the UK, including what strength means, who higher-strength pouches suit and where to shop them at Kangoo.",
        "secondary": ["strong nicotine pouches", "extra strong nicotine pouches", "strongest snus", "pablo nicotine pouches", "killa nicotine pouches"],
        "intent": "Commercial investigation",
        "links": [
            ("strong nicotine pouches", LINKS["strong"]),
            ("extra strong nicotine pouches", LINKS["extra"]),
            ("PABLO nicotine pouches", LINKS["pablo"]),
            ("KILLA nicotine pouches", LINKS["killa"]),
            ("strength ladder", LINKS["strength_ladder"]),
        ],
        "summary": "Capture strongest/strong snus demand while keeping claims careful and directing experienced users to strength pages.",
        "sections": [
            ("Quick answer", [
                "The strongest nicotine pouches are usually the products with the highest mg rating or most intense pouch feel, but strength should not be treated as a leaderboard for everyone.",
                "Strong and extra-strong pouches are best positioned for experienced adult nicotine users who already know they prefer a more powerful product.",
            ]),
            ("What makes a pouch feel strong", [
                "The mg number matters, but it is not the only factor. Moisture, pouch size, pH, flavour cooling and how long the pouch is used can all affect perceived intensity.",
                "Mint and ice flavours can sometimes feel sharper, while fruit and coffee flavours may feel softer even at similar strengths.",
            ]),
            ("Brands to compare", [
                "PABLO and KILLA are useful brand pages for shoppers looking at higher-strength products on Kangoo. ZYN and VELO also include products that can sit in strong or extra-strong ranges depending on the specific can.",
                "Use brand pages to compare flavours, then use strength pages to narrow by intensity.",
            ]),
            ("How to shop strong pouches responsibly", [
                "Do not choose the strongest product just because it looks like better value. A high-strength pouch that does not suit your tolerance is not a better buy.",
                "A better method is to compare strong products first, then move to extra strong only if you already know the stronger pouch feel is what you want.",
            ]),
        ],
        "faq": [
            ("Are extra strong pouches for everyone?", "No. They are better suited to experienced adult nicotine users."),
            ("Is PABLO stronger than ZYN?", "It depends on the specific product and strength. Compare the live product pages rather than relying on brand alone."),
            ("What is the best strong pouch flavour?", "Mint and ice are common for strong pouches, but grape, berry and coffee-style flavours are also popular."),
        ],
    },
    {
        "title": "How To Use Nicotine Pouches: Placement, Timing and Tips",
        "focus": "how to use nicotine pouches",
        "meta_title": "How To Use Nicotine Pouches | Placement and Timing Guide",
        "meta_description": "Learn how adult nicotine users typically use nicotine pouches, including placement, timing, disposal, strength choice and common first-order tips.",
        "secondary": ["what are nicotine pouches", "nicotine pouch", "nicotine pouches side effects", "how long do nicotine pouches last", "nicotine pouches uk"],
        "intent": "Informational",
        "links": [
            ("pouch finder", LINKS["finder"]),
            ("strength ladder", LINKS["strength_ladder"]),
            ("shop nicotine pouches", LINKS["all"]),
            ("99p trial pouches", LINKS["trial"]),
            ("mint nicotine pouches", LINKS["mint"]),
        ],
        "summary": "A careful usage explainer for adults, with no cessation or health claims.",
        "sections": [
            ("Quick answer", [
                "Nicotine pouches are typically placed under the upper lip, between the lip and gum. The pouch is not chewed or swallowed. After use, it should be removed and disposed of responsibly.",
                "Always follow the product packaging and choose a strength that matches your experience level. If you are unsure, start by comparing lower or balanced strengths before moving higher.",
            ]),
            ("Step-by-step use", [
                "First, place one pouch under your upper lip. Let it sit comfortably rather than moving it around constantly.",
                "Second, leave it in place for the time recommended on the product packaging or until the flavour and nicotine release feel finished.",
                "Third, remove the pouch and put it in a bin. Do not flush pouches, leave them where children or pets can reach them, or reuse a pouch.",
            ]),
            ("Choosing your first strength", [
                "If you are comparing pouches for the first time, the strength ladder is a better starting point than guessing from the strongest product available.",
                "Lower strengths suit a gentler feel. Strong and extra-strong pouches should be kept for experienced adult nicotine users.",
            ]),
            ("Common mistakes to avoid", [
                "Do not use multiple pouches at once to test a product. Do not chew or swallow the pouch. Do not choose extra strong products if you are unsure of your tolerance.",
                "If a pouch feels uncomfortable, remove it. This article is not medical advice, and anyone with health concerns should speak to a healthcare professional.",
            ]),
        ],
        "faq": [
            ("Do you chew nicotine pouches?", "No. They are normally placed under the lip and then removed after use."),
            ("Can you swallow nicotine pouches?", "No. Remove and dispose of the pouch after use."),
            ("How long should you use one pouch?", "Follow the product packaging. Time can vary by product and personal preference."),
        ],
    },
    {
        "title": "What Is Snus? Meaning, UK Rules and Tobacco-Free Alternatives",
        "focus": "what is snus",
        "meta_title": "What Is Snus? UK Meaning, Rules and Alternatives",
        "meta_description": "Understand what snus means, how tobacco snus differs from nicotine pouches, and why UK shoppers often search for tobacco-free pouch alternatives.",
        "secondary": ["snus meaning", "snus uk", "pablo snus", "velo snus", "nicotine pouches"],
        "intent": "Informational",
        "links": [
            ("shop nicotine pouches", LINKS["all"]),
            ("PABLO nicotine pouches", LINKS["pablo"]),
            ("VELO nicotine pouches", LINKS["velo"]),
            ("ZYN nicotine pouches", LINKS["zyn"]),
            ("compare pouches", LINKS["compare"]),
        ],
        "summary": "Answer the huge snus query while clarifying UK rules and moving users toward legal tobacco-free nicotine pouch pages.",
        "sections": [
            ("Quick answer", [
                "Snus traditionally means a moist oral tobacco product placed between the lip and gum. In the UK, many shoppers use the word 'snus' casually when they are actually searching for tobacco-free nicotine pouches.",
                "That distinction matters. Tobacco snus and tobacco-free nicotine pouches are not the same product, and UK rules treat tobacco oral snuff differently.",
            ]),
            ("What snus means", [
                "Traditional snus contains tobacco. It is usually supplied loose or in small pouches and is used orally rather than smoked.",
                "Because the format looks similar to modern nicotine pouches, people often search for terms such as 'VELO snus', 'PABLO snus' or 'ZYN snus' even when the product they want is a tobacco-free nicotine pouch.",
            ]),
            ("Snus and UK rules", [
                "GOV.UK guidance describes oral snuff, also known as snus, as a tobacco product and states that the sale of oral snuff is prohibited in the UK.",
                "This is why a UK ecommerce page should be careful with wording. Use 'nicotine pouches' for tobacco-free products, and explain that shoppers may use the word snus informally.",
            ]),
            ("Tobacco-free alternatives", [
                "Kangoo sells tobacco-free nicotine pouch products for adults who already use nicotine. Popular brand pages include ZYN, VELO, PABLO and KILLA.",
                "A useful next step is to compare pouches by strength and flavour rather than by the word 'snus' alone.",
            ]),
        ],
        "faq": [
            ("Is snus the same as nicotine pouches?", "No. Traditional snus contains tobacco, while modern nicotine pouches are usually tobacco-free."),
            ("Can you buy snus in the UK?", "GOV.UK guidance says the sale of oral snuff is prohibited in the UK."),
            ("Why do people say ZYN snus or VELO snus?", "Many shoppers use snus as a shorthand for oral nicotine pouches, even though the products are different."),
        ],
    },
    {
        "title": "Snus vs Nicotine Pouches: What Is the Difference in the UK?",
        "focus": "snus vs nicotine pouches",
        "meta_title": "Snus vs Nicotine Pouches | UK Difference Explained",
        "meta_description": "Compare snus and nicotine pouches in the UK, including tobacco content, legal wording, strengths, flavours and where to shop tobacco-free pouches.",
        "secondary": ["snus uk", "nicotine pouches", "tobacco free snus", "velo snus", "zyn snus"],
        "intent": "Informational and commercial",
        "links": [
            ("shop nicotine pouches", LINKS["all"]),
            ("tobacco-free ZYN pouches", LINKS["zyn"]),
            ("VELO nicotine pouches", LINKS["velo"]),
            ("pouch comparison tool", LINKS["compare"]),
            ("strength ladder", LINKS["strength_ladder"]),
        ],
        "summary": "A side-by-side explainer that can rank for snus education while internally linking commercial pouch pages.",
        "sections": [
            ("Quick answer", [
                "The main difference is tobacco. Traditional snus is an oral tobacco product. Nicotine pouches are designed as tobacco-free pouches that contain nicotine, flavourings and pouch material instead of tobacco leaf.",
                "In the UK, this distinction is important because traditional oral snuff has sale restrictions, while tobacco-free nicotine pouch products are sold through adult-only retail channels.",
            ]),
            ("Snus in simple terms", [
                "Snus is a tobacco product that is placed between the lip and gum. It is associated with Sweden and is often searched in the UK by people who are learning about oral nicotine formats.",
                "Because the pouch format is familiar, people sometimes call all oral pouches 'snus'. That is not technically accurate.",
            ]),
            ("Nicotine pouches in simple terms", [
                "Nicotine pouches are tobacco-free pouches for adults who already use nicotine. They come in different strengths, formats and flavours, including mint, berry, citrus, coffee and ice-style profiles.",
                "On Kangoo, the best way to shop them is by brand, strength or flavour rather than by the word snus.",
            ]),
            ("How to choose between search terms", [
                "If you are writing for UK shoppers, use 'snus' only to explain meaning, search language or legal differences. Use 'nicotine pouches' for Kangoo products and product category pages.",
                "This keeps the content accurate and reduces the risk of confusing tobacco snus with tobacco-free pouch products.",
            ]),
        ],
        "faq": [
            ("Is tobacco-free snus the same as nicotine pouches?", "Many shoppers use that phrase, but nicotine pouches is the clearer term for tobacco-free products."),
            ("Why does Kangoo sell nicotine pouches, not tobacco snus?", "Kangoo focuses on tobacco-free nicotine pouches for adults."),
            ("Which page should I visit after reading this?", "Start with the nicotine pouch category or the comparison tool."),
        ],
    },
    {
        "title": "ZYN Nicotine Pouches Guide: Strengths, Flavours and Best Picks",
        "focus": "zyn nicotine pouches",
        "meta_title": "ZYN Nicotine Pouches UK | Strengths and Flavours",
        "meta_description": "Explore ZYN nicotine pouches at Kangoo, including popular strengths, mint and fruit flavours, mini formats and links to ZYN product picks.",
        "secondary": ["zyn pouches", "zyn uk", "zyn flavours", "zyn strengths", "zyn cool mint"],
        "intent": "Commercial investigation",
        "links": [
            ("shop ZYN nicotine pouches", LINKS["zyn"]),
            ("ZYN Cool Mint Mini 3mg", LINKS["zyn_cool"]),
            ("ZYN Black Cherry Mini 3mg", LINKS["zyn_cherry"]),
            ("ZYN Citrus Mini 6mg", LINKS["zyn_citrus"]),
            ("compare ZYN with other pouches", LINKS["compare"]),
        ],
        "summary": "Brand guide for ZYN searches with product and category links.",
        "sections": [
            ("Quick answer", [
                "ZYN nicotine pouches are a popular tobacco-free pouch brand searched for mint, fruit and mini-format options. At Kangoo, the ZYN category is the best place to compare live stock, strength and flavour.",
                "Use this guide to move from broad searches like 'ZYN pouches' or 'ZYN UK' into practical choices by strength and flavour.",
            ]),
            ("Popular ZYN flavour directions", [
                "Mint-led options such as Cool Mint and Spearmint are useful for shoppers who want a clean, fresh pouch profile.",
                "Fruit-led options such as Black Cherry, Citrus and Red Fruits are better for shoppers who want a sweeter or sharper flavour profile.",
            ]),
            ("ZYN strengths and formats", [
                "ZYN products can include lower-strength mini styles as well as stronger options depending on the live range. Always compare the exact product page rather than assuming every ZYN pouch feels the same.",
                "Mini pouches are often useful for shoppers who want a smaller format and a more discreet feel.",
            ]),
            ("How to choose ZYN at Kangoo", [
                "If you are unsure, compare one mint and one fruit option first. If you already know you like ZYN, use the ZYN category page to check pack pricing and stock before adding products to basket.",
                "Internal links should point to the ZYN category first, then to individual products only when those products are in stock.",
            ]),
        ],
        "faq": [
            ("What ZYN flavour should I try first?", "Cool Mint, Spearmint, Citrus and Black Cherry are useful comparison points."),
            ("Are ZYN pouches tobacco-free?", "ZYN products sold as nicotine pouches are tobacco-free pouch products."),
            ("Where should ZYN links point?", "Use the ZYN category page for evergreen links, then add product links when live stock supports them."),
        ],
    },
    {
        "title": "VELO Nicotine Pouches Guide: Flavours, Strengths and Best Picks",
        "focus": "velo nicotine pouches",
        "meta_title": "VELO Nicotine Pouches UK | Flavours and Strengths",
        "meta_description": "Compare VELO nicotine pouches at Kangoo, including peppermint, spearmint, berry and stronger options with links to live VELO products.",
        "secondary": ["velo pouches", "velo snus", "velo flavours", "velo strengths", "velo freezing peppermint"],
        "intent": "Commercial investigation",
        "links": [
            ("shop VELO nicotine pouches", LINKS["velo"]),
            ("VELO Bright Spearmint 6mg", LINKS["velo_bright"]),
            ("VELO Freezing Peppermint 11mg", LINKS["velo_freezing"]),
            ("VELO Ruby Berry 10mg", LINKS["velo_ruby"]),
            ("compare VELO with other brands", LINKS["compare"]),
        ],
        "summary": "Brand guide for VELO searches and flavour terms.",
        "sections": [
            ("Quick answer", [
                "VELO nicotine pouches are a popular tobacco-free pouch range with strong search demand around mint, peppermint, berry and strength comparisons.",
                "The VELO category page is the best evergreen destination because stock and product mix can change over time.",
            ]),
            ("Popular VELO flavours", [
                "Mint and peppermint options are the most obvious starting point for shoppers who want a cooling pouch. Bright Spearmint, Crispy Peppermint and Freezing Peppermint-style searches all sit naturally in this cluster.",
                "Fruit options such as Ruby Berry, Tropical Mango, Orange Ice and Watermelon-style flavours help VELO capture shoppers who do not want a purely mint-led order.",
            ]),
            ("VELO strengths", [
                "VELO products can range from smoother everyday strengths to stronger products, depending on the exact can. Always compare the mg on the product page before buying.",
                "The strength ladder can help shoppers choose between balanced, strong and extra-strong products without relying on brand name alone.",
            ]),
            ("How to shop VELO at Kangoo", [
                "Use the VELO category page as the main place to browse current options, then check individual product pages for live stock, strength and price.",
                "Mint nicotine pouches, berry nicotine pouches and the comparison tool can help narrow the range if you are deciding between flavours.",
            ]),
        ],
        "faq": [
            ("What VELO flavour is most searched?", "Peppermint, spearmint and berry-led VELO searches are strong in the keyword exports."),
            ("Is VELO the same as snus?", "VELO nicotine pouches are tobacco-free pouches; snus traditionally refers to an oral tobacco product."),
            ("How do I compare VELO strengths?", "Use the product page mg and Kangoo's strength ladder."),
        ],
    },
    {
        "title": "ZYN vs VELO: Which Nicotine Pouch Brand Should You Try?",
        "focus": "zyn vs velo",
        "meta_title": "ZYN vs VELO | Nicotine Pouch Brand Comparison",
        "meta_description": "Compare ZYN and VELO nicotine pouches by flavour, strength, format and shopping style, with links to both Kangoo brand ranges.",
        "secondary": ["zyn nicotine pouches", "velo nicotine pouches", "best nicotine pouches", "compare nicotine pouches", "zyn flavours"],
        "intent": "Commercial investigation",
        "links": [
            ("shop ZYN nicotine pouches", LINKS["zyn"]),
            ("shop VELO nicotine pouches", LINKS["velo"]),
            ("compare pouches", LINKS["compare"]),
            ("mint nicotine pouches", LINKS["mint"]),
            ("berry nicotine pouches", LINKS["berry"]),
        ],
        "summary": "A comparison post that avoids unsupported winner claims and pushes to both brand pages.",
        "sections": [
            ("Quick answer", [
                "ZYN and VELO are both strong choices for adult shoppers comparing tobacco-free nicotine pouches. ZYN is useful for mini formats and clean flavour comparisons, while VELO is strong for broad flavour coverage across mint, berry, citrus and ice-style products.",
                "The better choice depends on the specific product, not just the brand name.",
            ]),
            ("Compare by flavour", [
                "If you want mint, compare ZYN Cool Mint or Spearmint-style products with VELO peppermint and spearmint products.",
                "If you want fruit, compare ZYN Black Cherry, Citrus or Red Fruits-style products with VELO Ruby Berry, Tropical Mango or Watermelon-style products.",
            ]),
            ("Compare by strength", [
                "Both brands can include multiple strengths. Check the exact mg and product format on each product page before adding it to basket.",
                "If strength is your main decision, use the strength ladder first and then choose the brand that has the best matching flavour in stock.",
            ]),
            ("Compare by shopping style", [
                "ZYN is a useful choice if you like mini pouches and classic flavour names. VELO is useful if you want a wider flavour exploration path.",
                "For WordPress internal linking, this post should link equally to both brand category pages and then to the comparison tool.",
            ]),
        ],
        "faq": [
            ("Is ZYN better than VELO?", "Not universally. The better choice depends on flavour, strength, format and stock."),
            ("Which is better for mint?", "Both brands have mint-led options. Compare exact products rather than brand alone."),
            ("Should this post name one winner?", "No. A balanced comparison is more durable and easier to keep accurate."),
        ],
    },
    {
        "title": "PABLO Nicotine Pouches Guide: Strong Pouches, Flavours and Who They Suit",
        "focus": "pablo nicotine pouches",
        "meta_title": "PABLO Nicotine Pouches UK | Strong Pouch Guide",
        "meta_description": "Explore PABLO nicotine pouches at Kangoo, including stronger pouch options, popular flavours and guidance for experienced adult users.",
        "secondary": ["pablo snus", "pablo pouches", "pablo 50mg", "pablo flavours", "pablo nicotine pouches uk"],
        "intent": "Commercial investigation",
        "links": [
            ("shop PABLO nicotine pouches", LINKS["pablo"]),
            ("PABLO Ice Cold 24mg", LINKS["pablo_ice"]),
            ("PABLO Grape Ice 30mg", LINKS["pablo_grape"]),
            ("extra strong nicotine pouches", LINKS["extra"]),
            ("strength ladder", LINKS["strength_ladder"]),
        ],
        "summary": "Capture PABLO/snus demand while keeping strong-product wording careful.",
        "sections": [
            ("Quick answer", [
                "PABLO nicotine pouches are best positioned for experienced adult nicotine users who are specifically comparing stronger pouch options.",
                "Many people search for 'PABLO snus', but for Kangoo content the clearer wording is PABLO nicotine pouches because the products are tobacco-free pouches.",
            ]),
            ("Why PABLO is searched with strong terms", [
                "The keyword exports show strong interest around PABLO, strength and snus-style wording. That means the guide should answer those searches while keeping the product language accurate.",
                "Do not present high strength as automatically better. Explain who it suits and link to the strength ladder for comparison.",
            ]),
            ("PABLO flavours to compare", [
                "Ice Cold-style products suit shoppers looking for a cooling profile. Grape Ice-style products add a fruit-led option while still sitting in a stronger pouch context.",
                "If a product is out of stock, link to the PABLO category page and extra strong category page rather than forcing a dead-end product link.",
            ]),
            ("How to buy PABLO at Kangoo", [
                "Use the PABLO category as the main evergreen internal link. Use individual product links only when stock is reliable.",
                "Add links to strong and extra strong nicotine pouch pages so shoppers can compare PABLO against other products in the same intensity band.",
            ]),
        ],
        "faq": [
            ("Is PABLO for beginners?", "PABLO is usually better suited to experienced adult nicotine users, especially in stronger formats."),
            ("Is PABLO snus?", "PABLO nicotine pouches are tobacco-free pouches; traditional snus refers to oral tobacco."),
            ("Where should PABLO content link?", "Use the PABLO category, extra strong category and strength ladder."),
        ],
    },
    {
        "title": "KILLA Nicotine Pouches Guide: Flavours, Strengths and Buying Tips",
        "focus": "killa nicotine pouches",
        "meta_title": "KILLA Nicotine Pouches UK | Flavours and Strengths",
        "meta_description": "Compare KILLA nicotine pouches at Kangoo, including flavour options, stronger pouch positioning and links to live KILLA products.",
        "secondary": ["killa snus", "killa nicopods", "killa nicotine pouches uk", "killa snus flavours", "strong nicotine pouches"],
        "intent": "Commercial investigation",
        "links": [
            ("shop KILLA nicotine pouches", LINKS["killa"]),
            ("strong nicotine pouches", LINKS["strong"]),
            ("extra strong nicotine pouches", LINKS["extra"]),
            ("flavour explorer", LINKS["flavour_explorer"]),
            ("compare pouches", LINKS["compare"]),
        ],
        "summary": "Brand guide for KILLA demand, bridging 'killa snus' language to nicotine pouch wording.",
        "sections": [
            ("Quick answer", [
                "KILLA nicotine pouches are searched by shoppers looking for bold flavours, stronger pouch options and snus-style terminology.",
                "For Kangoo content, use KILLA nicotine pouches as the main wording, then explain that some shoppers may search for KILLA snus even though tobacco-free pouches are a different product type.",
            ]),
            ("KILLA flavour positioning", [
                "KILLA can work well in content about bold flavour choice: mint, ice, fruit, coffee-style and sweet pouch directions all fit the brand's search profile.",
                "Use the flavour explorer to help shoppers compare flavour families before choosing a can.",
            ]),
            ("KILLA and strength", [
                "KILLA is often associated with stronger products. That makes it a good internal link from strong and extra-strong nicotine pouch pages.",
                "The content should be careful not to encourage users to jump straight to the highest strength. Strength should match experience level.",
            ]),
            ("How to shop KILLA at Kangoo", [
                "The KILLA category page should be the main destination for this post. Add supporting links to strong, extra strong and compare pouches pages.",
                "When products rotate, the category page keeps the article useful even if individual products go out of stock.",
            ]),
        ],
        "faq": [
            ("Is KILLA a strong nicotine pouch brand?", "Many KILLA products are positioned around stronger pouch experiences, but check the exact product page."),
            ("Is KILLA snus?", "KILLA nicotine pouches are tobacco-free pouches; traditional snus is a tobacco product."),
            ("What KILLA flavour should I compare first?", "Start with mint or ice if you want freshness, or fruit/coffee-style profiles if you want something different."),
        ],
    },
    {
        "title": "Mint Nicotine Pouches Guide: Spearmint, Peppermint and Ice Options",
        "focus": "mint nicotine pouches",
        "meta_title": "Mint Nicotine Pouches UK | Spearmint, Peppermint and Ice",
        "meta_description": "Compare mint nicotine pouches at Kangoo, including spearmint, peppermint, cool mint, ice flavours and links to ZYN and VELO picks.",
        "secondary": ["spearmint nicotine pouches", "peppermint nicotine pouches", "zyn cool mint", "velo freezing peppermint", "mint snus"],
        "intent": "Commercial investigation",
        "links": [
            ("mint nicotine pouches", LINKS["mint"]),
            ("VELO Bright Spearmint 6mg", LINKS["velo_bright"]),
            ("VELO Freezing Peppermint 11mg", LINKS["velo_freezing"]),
            ("ZYN Cool Mint Mini 3mg", LINKS["zyn_cool"]),
            ("flavour explorer", LINKS["flavour_explorer"]),
        ],
        "summary": "Flavour guide built around the strongest flavour keywords in the exports.",
        "sections": [
            ("Quick answer", [
                "Mint nicotine pouches are one of the easiest flavour families to compare because they cover spearmint, peppermint, cool mint, menthol and ice-style profiles.",
                "For Kangoo shoppers, mint is a useful everyday starting point before trying fruit, berry or stronger ice-led products.",
            ]),
            ("Spearmint vs peppermint", [
                "Spearmint usually feels smoother and sweeter. Peppermint is often sharper and cooler.",
                "A good internal comparison is VELO Bright Spearmint for a smoother mint direction and VELO Freezing Peppermint for a stronger cooling direction.",
            ]),
            ("Cool mint and ice-style pouches", [
                "Cool mint and ice-style names usually indicate a fresher, colder flavour profile. These can feel more intense even when the mg is similar to a non-mint product.",
                "That is why shoppers should compare flavour and strength together rather than choosing by name alone.",
            ]),
            ("Best internal links for mint content", [
                "Use the mint nicotine pouches page as the main commercial destination, then support it with ZYN, VELO and flavour explorer links.",
                "If a specific product has strong stock, include it in the product picks section of the WordPress post.",
            ]),
        ],
        "faq": [
            ("What is the difference between spearmint and peppermint pouches?", "Spearmint is usually smoother; peppermint is usually sharper and cooler."),
            ("Are mint pouches stronger?", "Not automatically. Mint can feel cooling, but strength depends on the product."),
            ("Which brands sell mint nicotine pouches?", "ZYN and VELO are useful starting points on Kangoo."),
        ],
    },
    {
        "title": "Berry Nicotine Pouches Guide: Cherry, Berry and Fruit Pouch Picks",
        "focus": "berry nicotine pouches",
        "meta_title": "Berry Nicotine Pouches UK | Cherry and Fruit Pouch Guide",
        "meta_description": "Compare berry nicotine pouches at Kangoo, including cherry, black cherry, ruby berry and fruit-led pouch options from ZYN and VELO.",
        "secondary": ["cherry nicotine pouches", "black cherry nicotine pouches", "velo ruby berry", "zyn black cherry", "fruit nicotine pouches"],
        "intent": "Commercial investigation",
        "links": [
            ("berry nicotine pouches", LINKS["berry"]),
            ("ZYN Black Cherry Mini 3mg", LINKS["zyn_cherry"]),
            ("VELO Ruby Berry 10mg", LINKS["velo_ruby"]),
            ("flavour explorer", LINKS["flavour_explorer"]),
            ("shop nicotine pouches", LINKS["all"]),
        ],
        "summary": "A flavour article for fruit/berry search demand and product links.",
        "sections": [
            ("Quick answer", [
                "Berry nicotine pouches are a good option for shoppers who want a fruit-led pouch rather than a mint-heavy flavour.",
                "Common berry-style searches include cherry, black cherry, ruby berry, red fruits and mixed fruit pouch names.",
            ]),
            ("Cherry and black cherry pouches", [
                "Cherry and black cherry pouches tend to feel richer and sweeter than citrus or mint options.",
                "ZYN Black Cherry-style products are useful product links when stock is available because they connect a clear flavour query with a recognizable brand.",
            ]),
            ("Berry vs citrus vs tropical", [
                "Berry flavours are usually softer and sweeter. Citrus flavours are sharper and brighter. Tropical flavours can feel warmer and juicier.",
                "The flavour explorer can help shoppers move from broad flavour preference into specific products.",
            ]),
            ("How to shop berry pouches at Kangoo", [
                "Use the berry nicotine pouches page as the evergreen destination. Then link to ZYN and VELO product examples when stock supports it.",
                "Add a comparison link so shoppers can check strength and price before choosing a regular can.",
            ]),
        ],
        "faq": [
            ("Are berry pouches only low strength?", "No. Berry flavours can appear across different strengths depending on the product."),
            ("Which brands have berry pouches?", "ZYN and VELO are useful starting points, and trial stock can rotate."),
            ("Are berry pouches sweeter than mint?", "Usually, yes. Berry profiles are generally more fruit-led than cooling mint profiles."),
        ],
    },
    {
        "title": "Are Nicotine Pouches Legal in the UK? Snus, Age Rules and Buying Online",
        "focus": "are nicotine pouches legal in the uk",
        "meta_title": "Are Nicotine Pouches Legal in the UK? Rules Explained",
        "meta_description": "Learn the difference between UK nicotine pouch shopping, tobacco snus restrictions, adult-only positioning and what to check before buying online.",
        "secondary": ["snus uk", "is snus legal in uk", "nicotine pouches uk", "tobacco free nicotine pouches", "buy nicotine pouches online"],
        "intent": "Informational",
        "links": [
            ("shop nicotine pouches", LINKS["all"]),
            ("what is snus guide destination", LINKS["all"]),
            ("compare pouches", LINKS["compare"]),
            ("ZYN nicotine pouches", LINKS["zyn"]),
            ("VELO nicotine pouches", LINKS["velo"]),
        ],
        "summary": "Legal explainer, grounded in official GOV language, useful for snus confusion and trust.",
        "sections": [
            ("Quick answer", [
                "Tobacco-free nicotine pouches are sold in the UK through adult-only retail channels. Traditional tobacco snus is different and should not be confused with nicotine pouches.",
                "GOV.UK guidance states that oral snuff, also known as snus, is a tobacco product and that the sale of oral snuff is prohibited in the UK.",
            ]),
            ("Nicotine pouches vs tobacco snus", [
                "Nicotine pouches are tobacco-free products for adults who already use nicotine. Traditional snus contains tobacco and sits under a different legal and product definition.",
                "This is why Kangoo content should use nicotine pouches as the main product wording and use snus only when explaining search language or legal differences.",
            ]),
            ("Buying online", [
                "When buying online, shoppers should check the retailer, product strength, delivery terms, stock status and age-gating information.",
                "Kangoo supports adult-only positioning, discreet packaging and category pages that help shoppers compare products by brand, strength and flavour.",
            ]),
            ("What not to claim", [
                "Do not describe nicotine pouches as medical products or licensed stop-smoking products. NHS guidance says nicotine pouches are not licensed stop-smoking products in the UK.",
                "Keep product copy factual: brand, flavour, strength, price and delivery terms are safer content areas than unsupported health claims.",
            ]),
        ],
        "faq": [
            ("Is snus legal to sell in the UK?", "GOV.UK guidance states that the sale of oral snuff is prohibited in the UK."),
            ("Are nicotine pouches the same as oral snuff?", "No. Tobacco-free nicotine pouches are different from traditional tobacco oral snuff."),
            ("Can Kangoo say nicotine pouches help stop smoking?", "No. The safer wording is that they are for adults who already use nicotine products."),
        ],
    },
    {
        "title": "Tobacco-Free Nicotine Pouches Explained: What Is Inside a Pouch?",
        "focus": "tobacco free nicotine pouches",
        "meta_title": "Tobacco-Free Nicotine Pouches Explained | Kangoo Guide",
        "meta_description": "Understand tobacco-free nicotine pouches, how they differ from snus, what shoppers compare and where to browse brands like ZYN, VELO, PABLO and KILLA.",
        "secondary": ["nicotine pouches", "tobacco free snus", "what are nicotine pouches", "oral nicotine products", "nicotine pouch"],
        "intent": "Informational",
        "links": [
            ("shop tobacco-free nicotine pouches", LINKS["all"]),
            ("ZYN nicotine pouches", LINKS["zyn"]),
            ("VELO nicotine pouches", LINKS["velo"]),
            ("PABLO nicotine pouches", LINKS["pablo"]),
            ("KILLA nicotine pouches", LINKS["killa"]),
        ],
        "summary": "Educational page explaining tobacco-free pouches without drifting into medical claims.",
        "sections": [
            ("Quick answer", [
                "Tobacco-free nicotine pouches are small oral pouches that contain nicotine but do not contain tobacco leaf. They are placed under the lip and removed after use.",
                "They are different from traditional snus, which contains tobacco.",
            ]),
            ("What is usually inside a pouch", [
                "Product formulas vary by brand, but tobacco-free pouches commonly include nicotine, plant-based pouch material, flavourings, fillers and moisture-control ingredients.",
                "The exact product page and packaging should always be treated as the source of truth.",
            ]),
            ("How shoppers compare tobacco-free pouches", [
                "The most useful comparison points are brand, strength, flavour, pouch size, price and stock.",
                "A shopper who wants freshness might start with mint. A shopper who wants fruit might compare berry or citrus. A shopper who wants higher intensity should use strength pages carefully.",
            ]),
            ("Best Kangoo pages to use next", [
                "Use the main nicotine pouches category for the full range, then move into brand pages when you know the brand you want.",
                "The comparison tool and pouch finder help turn broad education into a buying decision.",
            ]),
        ],
        "faq": [
            ("Do tobacco-free nicotine pouches contain tobacco?", "No. Tobacco-free pouches do not contain tobacco leaf."),
            ("Are nicotine pouches the same as snus?", "No. Traditional snus contains tobacco; nicotine pouches are tobacco-free."),
            ("What should I compare before buying?", "Compare strength, flavour, brand, price, pouch count and stock."),
        ],
    },
    {
        "title": "Nicotine Pouches vs Vapes: Format Differences for Adult Nicotine Users",
        "focus": "nicotine pouches vs vapes",
        "meta_title": "Nicotine Pouches vs Vapes | Key Format Differences",
        "meta_description": "Compare nicotine pouches and vapes by format, flavour, discretion, strength choices and shopping considerations for adult nicotine users.",
        "secondary": ["alternatives to e cigarettes", "nicotine pouches uk", "what are nicotine pouches", "smokeless nicotine", "oral nicotine products"],
        "intent": "Informational",
        "links": [
            ("shop nicotine pouches", LINKS["all"]),
            ("pouch finder", LINKS["finder"]),
            ("compare pouches", LINKS["compare"]),
            ("strength ladder", LINKS["strength_ladder"]),
            ("99p pouch trials", LINKS["trial"]),
        ],
        "summary": "Informational comparison that avoids claiming pouches are a cessation product or healthier option.",
        "sections": [
            ("Quick answer", [
                "Nicotine pouches and vapes are different formats. A nicotine pouch is placed under the lip and removed after use. A vape is inhaled through a device.",
                "This article should compare practical format differences only, not make medical or stop-smoking claims.",
            ]),
            ("Format and use", [
                "Nicotine pouches do not require a device, charging, e-liquid or inhalation. They are chosen by strength, flavour and pouch format.",
                "Vapes involve a device, vapour and inhalation. The shopping considerations are different, so comparisons should stay factual and avoid saying one format is universally better.",
            ]),
            ("Flavour and strength", [
                "Both categories use flavour language, but pouch flavour is experienced orally under the lip rather than through vapour.",
                "Pouch strength is usually compared by mg and perceived pouch feel. For Kangoo shoppers, the strength ladder is the cleanest way to understand the range.",
            ]),
            ("Where to go from here", [
                "If a shopper wants to explore pouches, send them to the 99p trial range, the pouch finder and the main nicotine pouches category.",
                "If they want product-by-product comparison, the comparison tool is the strongest internal link.",
            ]),
        ],
        "faq": [
            ("Are nicotine pouches vapes?", "No. They are different formats."),
            ("Do nicotine pouches need a device?", "No. A pouch is placed under the lip and removed after use."),
            ("Can this article say pouches help stop vaping?", "No. Keep the article to factual format differences only."),
        ],
    },
]

INITIAL_PUBLISH_START = date(2026, 5, 21)
EXTRA_PUBLISH_START = date(2026, 6, 8)


def make_extra_guide(
    title: str,
    focus: str,
    meta_title: str,
    meta_description: str,
    secondary: list[str],
    intent: str,
    links: list[tuple[str, str]],
    quick_answer: str,
    comparison_note: str,
    kangoo_route: str,
    care_note: str,
    faq: list[tuple[str, str]],
) -> dict:
    return {
        "title": title,
        "focus": focus,
        "meta_title": meta_title,
        "meta_description": meta_description,
        "secondary": secondary,
        "intent": intent,
        "links": links,
        "summary": quick_answer,
        "sections": [
            ("Quick answer", [
                quick_answer,
                "The useful shopping route is to compare flavour, strength, format, live stock and price together rather than choosing from one headline claim.",
            ]),
            ("What to compare before buying", [
                comparison_note,
                "Live category pages are the best place to check current options, while product pages show the exact strength, price, flavour and availability before you order.",
            ]),
            ("Best Kangoo route", [
                kangoo_route,
                "If you are unsure, the pouch finder, strength ladder and flavour explorer can help narrow the choice by experience level, strength and flavour.",
            ]),
            ("Important notes", [
                "Product details, stock, pricing and delivery terms can change, so the live product page is the source of truth before ordering.",
                "Nicotine is addictive. This guide is for adult nicotine users only and is not medical advice.",
            ]),
        ],
        "faq": faq,
    }


EXTRA_GUIDES = [
    make_extra_guide(
        "What Are Nicotine Pouches? UK Beginner Guide",
        "what are nicotine pouches",
        "What Are Nicotine Pouches? UK Beginner Guide | Kangoo",
        "Learn what nicotine pouches are, how they differ from tobacco snus and how adult nicotine users can compare flavours, strengths and brands.",
        ["nicotine pouch", "nicotine pouches uk", "what is a nicotine pouch", "tobacco free nicotine pouches", "how to use nicotine pouches"],
        "Informational",
        [("shop nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"]), ("flavour explorer", LINKS["flavour_explorer"])],
        "Nicotine pouches are small tobacco-free pouches designed to sit under the lip and deliver nicotine without smoke or vapour.",
        "Beginners should compare pouch strength, flavour, brand, pouch size and price before deciding what to try first.",
        "Start with the full nicotine pouch range, then narrow with the pouch finder if strength or flavour choice feels unclear.",
        "Nicotine is addictive, and nicotine pouches are not licensed stop-smoking products in the UK.",
        [("Are nicotine pouches tobacco-free?", "Most modern nicotine pouches sold by Kangoo are tobacco-free, but always check the individual product page."), ("Do you smoke nicotine pouches?", "No. They are placed under the lip and are not smoked, vaped, chewed or swallowed."), ("What strength should beginners compare first?", "Lower or balanced strengths are usually easier to compare before moving into strong products.")],
    ),
    make_extra_guide(
        "Nicotine Pouches Side Effects: What Adult Users Should Know",
        "nicotine pouches side effects",
        "Nicotine Pouches Side Effects | Adult User Guide",
        "A careful guide to possible nicotine pouch side effects, strength choice, pouch comfort and when adult users should stop using a pouch.",
        ["nicotine pouch side effects", "are nicotine pouches safe", "what do nicotine pouches do", "nicotine pouches gum", "strong nicotine pouches"],
        "Informational",
        [("strength ladder", LINKS["strength_ladder"]), ("pouch finder", LINKS["finder"]), ("shop nicotine pouches", LINKS["all"]), ("strong nicotine pouches", LINKS["strong"])],
        "Possible nicotine pouch side effects can include mouth irritation, hiccups, nausea, dizziness or a stronger-than-expected nicotine sensation, especially if the strength is too high.",
        "Strength, pouch size, moisture and use time can all change the experience, so shoppers should not compare by mg number alone.",
        "Use the strength ladder before choosing strong products, and choose lower or balanced options if you are unsure.",
        "This guide is not medical advice. Anyone with symptoms or health concerns should stop using the pouch and speak to a healthcare professional.",
        [("Can a pouch feel too strong?", "Yes. If a pouch feels uncomfortable, remove it and avoid using another high-strength pouch."), ("Are side effects the same for everyone?", "No. Individual tolerance and the product strength can change the experience."), ("Should I use extra strong pouches first?", "No. Extra strong pouches are better suited to experienced adult nicotine users.")],
    ),
    make_extra_guide(
        "Snus Meaning: What People Mean by Snus in the UK",
        "snus meaning",
        "Snus Meaning UK | Snus vs Nicotine Pouches",
        "Understand the meaning of snus in the UK, why people use the word for nicotine pouches and how tobacco-free pouches differ.",
        ["what is snus", "snus uk", "snus pouch", "snus nicotine", "tobacco free snus"],
        "Informational",
        [("tobacco-free nicotine pouches", LINKS["all"]), ("snus vs nicotine pouches guide", LINKS["all"]), ("ZYN nicotine pouches", LINKS["zyn"]), ("VELO nicotine pouches", LINKS["velo"])],
        "In the UK, many shoppers use the word snus loosely when they are actually looking for tobacco-free nicotine pouches.",
        "Traditional snus contains tobacco, while modern nicotine pouches are usually tobacco-free and should be compared as a separate product category.",
        "For Kangoo shoppers, the practical route is to browse tobacco-free nicotine pouches and use brand, flavour and strength pages instead of relying on the word snus.",
        "Be careful with legal wording: traditional oral snuff or snus is treated differently from tobacco-free nicotine pouches in the UK.",
        [("Is snus the same as nicotine pouches?", "No. People often use the terms together, but traditional snus contains tobacco and nicotine pouches are usually tobacco-free."), ("Why do UK shoppers search for snus?", "Many use snus as a shortcut phrase for under-lip nicotine products."), ("What should I buy if I searched for snus?", "Compare tobacco-free nicotine pouches by strength, flavour and brand.")],
    ),
    make_extra_guide(
        "Snus UK: Legal Alternatives to Traditional Snus",
        "snus uk",
        "Snus UK Guide | Legal Tobacco-Free Alternatives",
        "Learn why UK shoppers search for snus, how tobacco-free nicotine pouches differ and where to compare legal pouch alternatives.",
        ["snus legal uk", "buy snus uk", "tobacco free snus", "nicotine pouches uk", "snus pouches"],
        "Informational and commercial",
        [("shop nicotine pouches", LINKS["all"]), ("ZYN nicotine pouches", LINKS["zyn"]), ("VELO nicotine pouches", LINKS["velo"]), ("pouch finder", LINKS["finder"])],
        "UK searches for snus often point to tobacco-free nicotine pouches, because many shoppers use the word snus for pouch-style nicotine products.",
        "The key comparison is tobacco content, legal status, strength, flavour and format.",
        "Use Kangoo's nicotine pouch category for legal tobacco-free options, then narrow by ZYN, VELO, PABLO or KILLA if you already prefer a brand.",
        "Avoid implying traditional snus and tobacco-free nicotine pouches are identical. They are different product types.",
        [("Can you buy traditional snus in the UK?", "Traditional oral tobacco snus is prohibited for sale in the UK."), ("Are nicotine pouches the same legal category?", "No. Tobacco-free nicotine pouches are a different product type."), ("Why does Kangoo sell pouches, not tobacco snus?", "Kangoo focuses on tobacco-free nicotine pouches for adult nicotine users.")],
    ),
    make_extra_guide(
        "Nicotine Pouches vs Nicotine Gum, Lozenges and Patches",
        "nicotine pouches vs gum",
        "Nicotine Pouches vs Gum, Lozenges and Patches",
        "Compare nicotine pouches with nicotine gum, lozenges and patches by format, use style and shopping considerations for adult nicotine users.",
        ["nicotine gum", "nicotine lozenges", "nicotine patches", "nicotine replacement therapy", "nicotine pouch"],
        "Informational",
        [("shop nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"]), ("how to use nicotine pouches", LINKS["all"])],
        "Nicotine pouches, gum, lozenges and patches are different formats, and they should not be treated as interchangeable products.",
        "Gum, lozenges and patches can be licensed stop-smoking products, while nicotine pouches are not licensed stop-smoking products in the UK.",
        "Kangoo content should compare format and shopping factors, then send users to product pages only for tobacco-free pouch browsing.",
        "Do not position nicotine pouches as a medicine, treatment or guaranteed alternative to stop-smoking support.",
        [("Are nicotine pouches NRT?", "No. Nicotine pouches are not licensed stop-smoking products in the UK."), ("Are patches used the same way as pouches?", "No. Patches are worn on the skin, while pouches sit under the lip."), ("Should this guide replace medical advice?", "No. It is a format comparison, not medical advice.")],
    ),
    make_extra_guide(
        "ZYN Pouches Guide: Flavours, Strengths and Buying Tips",
        "zyn pouches",
        "ZYN Pouches UK | Flavours, Strengths and Buying Tips",
        "Compare ZYN pouches by flavour, strength and format, including mint, citrus and berry-style options available through Kangoo.",
        ["zyn", "zyns", "zyn nicotine", "zyn nicotine pouches", "zyn flavours"],
        "Commercial investigation",
        [("ZYN nicotine pouches", LINKS["zyn"]), ("ZYN Cool Mint", LINKS["zyn_cool"]), ("ZYN Black Cherry", LINKS["zyn_cherry"]), ("ZYN Citrus", LINKS["zyn_citrus"]), ("compare pouches", LINKS["compare"])],
        "ZYN pouches are one of the most searched pouch brands, with shoppers often comparing mini formats, mint flavours and fruit options.",
        "The key ZYN comparison points are strength, flavour family, pouch format and whether the specific can is in stock.",
        "Send shoppers to the ZYN category first, then use product links for reliable picks such as Cool Mint, Black Cherry and Citrus when live stock supports it.",
        "Avoid saying every ZYN product suits every user. Match recommendations to flavour and strength preference.",
        [("What ZYN flavour should I try first?", "Mint is a common starting point, while Black Cherry and Citrus suit fruit-led preferences."), ("Are all ZYN pouches the same strength?", "No. Check the exact mg on each product page."), ("Where should I compare ZYN with other brands?", "Use Kangoo's compare pouches page or the full nicotine pouch category.")],
    ),
    make_extra_guide(
        "What Is ZYN? UK Guide to ZYN Nicotine Pouches",
        "what is zyn",
        "What Is ZYN? UK Guide to ZYN Nicotine Pouches",
        "Find out what ZYN is, how ZYN nicotine pouches work and how to compare ZYN flavours and strengths at Kangoo.",
        ["zyn nicotine pouches", "zyns", "zyn pouches", "zyn nicotine", "what are zyns"],
        "Informational and commercial",
        [("ZYN nicotine pouches", LINKS["zyn"]), ("ZYN Cool Mint", LINKS["zyn_cool"]), ("ZYN Black Cherry", LINKS["zyn_cherry"]), ("shop nicotine pouches", LINKS["all"])],
        "ZYN is a tobacco-free nicotine pouch brand sold in multiple flavours and strengths for adult nicotine users.",
        "Shoppers should compare ZYN by exact strength, flavour profile, pouch format and stock rather than choosing by brand name alone.",
        "Use the ZYN category as the main evergreen link, then add product links for specific flavours when stock is live.",
        "Keep copy factual and avoid lifestyle or health claims around ZYN.",
        [("Is ZYN tobacco-free?", "ZYN nicotine pouches are generally sold as tobacco-free pouches, but always check the product page."), ("Is ZYN snus?", "Many people call it snus, but ZYN is a nicotine pouch rather than traditional tobacco snus."), ("Which ZYN should I compare first?", "Cool Mint, Black Cherry and Citrus are useful starting flavour families.")],
    ),
    make_extra_guide(
        "VELO Pouches Guide: Mint, Berry and Citrus Options",
        "velo pouches",
        "VELO Pouches UK | Mint, Berry and Citrus Guide",
        "Compare VELO pouches by flavour family, strength and pouch style, with links to popular VELO options at Kangoo.",
        ["velo", "velo nicotine", "velo nicotine pouches", "velo snus", "velo flavours"],
        "Commercial investigation",
        [("VELO nicotine pouches", LINKS["velo"]), ("VELO Bright Spearmint", LINKS["velo_bright"]), ("VELO Freezing Peppermint", LINKS["velo_freezing"]), ("VELO Ruby Berry", LINKS["velo_ruby"]), ("flavour explorer", LINKS["flavour_explorer"])],
        "VELO pouches are often compared for fresh mint, bright fruit and stronger cooling profiles.",
        "The useful comparison is flavour first, then strength and format, because VELO shoppers often know the flavour family they want.",
        "Use the VELO category as the hub, then link to flavour-led products such as Bright Spearmint, Freezing Peppermint and Ruby Berry when stocked.",
        "Do not describe VELO as risk-free or suitable for everyone; keep the article for adult nicotine users.",
        [("What VELO flavour is popular?", "Mint and peppermint profiles are common everyday comparisons, while berry suits fruit-led shoppers."), ("Is VELO the same as snus?", "VELO is usually searched alongside snus, but it is a tobacco-free nicotine pouch brand."), ("Where can I compare VELO strengths?", "Use the VELO category and Kangoo strength ladder.")],
    ),
    make_extra_guide(
        "Nordic Spirit Alternatives: How to Compare Nicotine Pouch Brands",
        "nordic spirit alternatives",
        "Nordic Spirit Alternatives | Compare Pouch Brands",
        "Looking for Nordic Spirit alternatives? Compare ZYN, VELO, PABLO and KILLA by flavour, strength and buying route at Kangoo.",
        ["nordic spirit", "nordic spirit nicotine pouches", "nordic spirit uk", "nicotine pouch brands", "best nicotine pouches uk"],
        "Commercial investigation",
        [("shop nicotine pouches", LINKS["all"]), ("ZYN nicotine pouches", LINKS["zyn"]), ("VELO nicotine pouches", LINKS["velo"]), ("compare pouches", LINKS["compare"])],
        "If a shopper knows Nordic Spirit but wants to compare other brands, the best route is to match flavour family and strength first.",
        "Compare mint, berry, citrus and strong profiles across brands rather than assuming one brand maps neatly to another.",
        "Kangoo should route this search to the full category, ZYN, VELO and comparison tools rather than making unsupported competitor claims.",
        "Keep the page framed as alternatives and comparison, not as an absolute claim that one brand is better.",
        [("Does Kangoo need to stock Nordic Spirit for this guide?", "No. The guide can help shoppers compare alternative pouch brands available at Kangoo."), ("What should Nordic Spirit shoppers compare first?", "Start with flavour family, then match strength and pouch format."), ("Can this guide name Nordic Spirit?", "Yes, if the copy stays factual and avoids unverifiable superiority claims.")],
    ),
    make_extra_guide(
        "PABLO Snus vs PABLO Nicotine Pouches: UK Buying Guide",
        "pablo snus",
        "PABLO Snus vs PABLO Nicotine Pouches | UK Guide",
        "Understand why UK shoppers search for PABLO snus and how to compare PABLO nicotine pouches by flavour and strength.",
        ["pablo nicotine pouches", "pablo pouches", "pablo ice cold", "pablo grape ice", "strong nicotine pouches"],
        "Commercial investigation",
        [("PABLO nicotine pouches", LINKS["pablo"]), ("PABLO Ice Cold", LINKS["pablo_ice"]), ("PABLO Grape Ice", LINKS["pablo_grape"]), ("extra strong nicotine pouches", LINKS["extra"])],
        "Many UK shoppers search for PABLO snus when they mean strong PABLO nicotine pouches.",
        "PABLO should be compared carefully because many products are positioned toward experienced users and higher strengths.",
        "Route shoppers to the PABLO brand page first, then to strong and extra-strong category pages if strength is their main filter.",
        "Avoid beginner-friendly wording around high-strength PABLO products.",
        [("Is PABLO snus?", "People often use that term, but Kangoo sells PABLO nicotine pouches rather than traditional tobacco snus."), ("Is PABLO strong?", "Many PABLO products are high-strength, so check each product page."), ("Who should compare PABLO?", "Experienced adult nicotine users who already prefer stronger pouches.")],
    ),
    make_extra_guide(
        "KILLA Snus vs KILLA Nicotine Pouches: Flavours and Strengths",
        "killa snus",
        "KILLA Snus vs KILLA Nicotine Pouches | Kangoo",
        "Compare KILLA nicotine pouches, why shoppers search for KILLA snus and which flavour and strength details matter before buying.",
        ["killa nicotine pouches", "killa pouches", "killa snus uk", "strong nicotine pouches", "nicotine pouch brands"],
        "Commercial investigation",
        [("KILLA nicotine pouches", LINKS["killa"]), ("strong nicotine pouches", LINKS["strong"]), ("extra strong nicotine pouches", LINKS["extra"]), ("pouch finder", LINKS["finder"])],
        "KILLA searches often use the word snus, but shoppers are usually comparing strong tobacco-free nicotine pouches.",
        "The useful comparison is flavour intensity, strength, pouch feel and whether the product suits an experienced user.",
        "Use KILLA as the brand hub and support it with strong and extra-strong landing pages.",
        "Keep the article clear that stronger does not mean better for every shopper.",
        [("Is KILLA a beginner pouch?", "Many KILLA products are better suited to experienced users, so check strength carefully."), ("Why do people say KILLA snus?", "Snus is often used loosely for pouch products in the UK."), ("Where should I compare KILLA?", "Start on the KILLA category page, then use strength pages if intensity matters most.")],
    ),
    make_extra_guide(
        "Why Do Footballers Use Snus? UK Nicotine Pouch Context",
        "footballers snus",
        "Why Do Footballers Use Snus? UK Pouch Context",
        "A careful UK guide to footballers, snus searches and nicotine pouches, with adult-user context and no performance claims.",
        ["snus footballers", "footballers nicotine pouches", "snus in football", "what does snus do", "snus side effects"],
        "Informational",
        [("what are nicotine pouches", LINKS["all"]), ("shop nicotine pouches", LINKS["all"]), ("strength ladder", LINKS["strength_ladder"]), ("pouch finder", LINKS["finder"])],
        "Searches around footballers and snus usually come from media coverage and curiosity about pouch-style nicotine products.",
        "This guide should explain terminology and product differences without suggesting nicotine improves sport performance.",
        "Route adult shoppers into education pages first, then into pouch finder and category pages only if they are already adult nicotine users.",
        "Do not make performance, focus or fitness claims for nicotine pouches.",
        [("Do nicotine pouches improve football performance?", "Do not rely on nicotine pouches for performance. This guide does not make performance claims."), ("Is this guide for under-18s?", "No. Kangoo content is for adult nicotine users only."), ("Why is snus mentioned in football?", "The term is often used in media discussions about pouch-style nicotine products.")],
    ),
    make_extra_guide(
        "Do Nicotine Pouches Stain Teeth? What to Know Before Buying",
        "do nicotine pouches stain teeth",
        "Do Nicotine Pouches Stain Teeth? UK Guide",
        "Learn what adult users should consider about nicotine pouches, teeth, flavour choice and pouch use without making unsupported safety claims.",
        ["nicotine pouches teeth", "nicotine pouches gum", "nicotine pouch side effects", "what do nicotine pouches do", "are nicotine pouches safe"],
        "Informational",
        [("shop nicotine pouches", LINKS["all"]), ("how to use nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("mint nicotine pouches", LINKS["mint"])],
        "Nicotine pouches do not contain tobacco leaf, but users may still want to understand mouth comfort, gum contact and flavour staining questions.",
        "The safest content approach is to explain format differences and tell users to seek dental advice for personal concerns.",
        "Use this guide to link to how-to content and lower-strength comparison routes rather than pushing strong products.",
        "Avoid saying nicotine pouches are safe for teeth or gums.",
        [("Can nicotine pouches affect my mouth?", "Some users may notice irritation or discomfort. Stop using the pouch if it feels uncomfortable."), ("Should I ask a dentist?", "Yes, for personal dental or gum concerns."), ("Are tobacco-free pouches the same as tobacco snus?", "No. They do not contain tobacco leaf, but they still contain nicotine.")],
    ),
    make_extra_guide(
        "How Long Do Nicotine Pouches Last? Use Time and Flavour Guide",
        "how long do nicotine pouches last",
        "How Long Do Nicotine Pouches Last? Use Time Guide",
        "Understand typical nicotine pouch use time, flavour fade, strength feel and when adult users should remove a pouch.",
        ["how to use nicotine pouches", "nicotine pouch", "what do nicotine pouches do", "nicotine pouches uk", "pouch strength"],
        "Informational",
        [("how to use nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"]), ("shop nicotine pouches", LINKS["all"])],
        "How long a nicotine pouch lasts depends on the product, flavour, moisture, strength and personal preference.",
        "A pouch should be removed when the product guidance says to remove it, or sooner if it feels uncomfortable.",
        "Use this guide to support how-to searches, then route shoppers toward strength and flavour pages.",
        "Do not encourage longer use than the product packaging recommends.",
        [("Should I chew a pouch to make it last longer?", "No. Nicotine pouches are not designed to be chewed."), ("Can flavour fade before the pouch is finished?", "Yes. Flavour fade and nicotine feel can vary by product."), ("What if a pouch feels uncomfortable?", "Remove it and avoid using another pouch that feels too strong.")],
    ),
    make_extra_guide(
        "Can You Swallow Nicotine Pouches? Placement and Disposal Guide",
        "can you swallow nicotine pouches",
        "Can You Swallow Nicotine Pouches? Placement Guide",
        "Learn why nicotine pouches are placed under the lip, not swallowed, and how adult users should remove and dispose of them.",
        ["how to use nicotine pouches", "do you chew nicotine pouches", "nicotine pouch", "what are nicotine pouches", "nicotine pouches side effects"],
        "Informational",
        [("how to use nicotine pouches", LINKS["all"]), ("shop nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"])],
        "Nicotine pouches are not meant to be swallowed. They are normally placed under the upper lip and removed after use.",
        "This search intent is practical, so the page should answer quickly before explaining placement, timing and disposal.",
        "Route readers to the how-to guide, then to the pouch finder if they are choosing a product.",
        "Keep disposal advice clear and avoid encouraging misuse.",
        [("Do you chew nicotine pouches?", "No. They are normally placed under the lip, not chewed."), ("What should I do after using one?", "Remove it and dispose of it responsibly in a bin."), ("What if I swallowed one accidentally?", "Seek appropriate medical advice or contact a healthcare professional if concerned.")],
    ),
    make_extra_guide(
        "Nicotine Pouch Flavours: Mint, Berry, Citrus, Coffee and Ice",
        "nicotine pouch flavours",
        "Nicotine Pouch Flavours | Mint, Berry, Citrus and Ice",
        "Compare nicotine pouch flavours including mint, berry, citrus, coffee and ice, with Kangoo links for adult users choosing a pouch.",
        ["nicotine pouches flavours", "mint nicotine pouches", "berry nicotine pouches", "citrus nicotine pouches", "coffee nicotine pouches"],
        "Commercial investigation",
        [("flavour explorer", LINKS["flavour_explorer"]), ("mint nicotine pouches", LINKS["mint"]), ("berry nicotine pouches", LINKS["berry"]), ("shop nicotine pouches", LINKS["all"])],
        "Flavour is often the easiest way to narrow a nicotine pouch order before comparing strength and brand.",
        "Mint and ice are common everyday choices, berry and citrus suit fruit-led shoppers, and coffee or tropical flavours add variety.",
        "Use the flavour explorer as the hub, then point to mint and berry landing pages plus the full category.",
        "Do not imply flavour changes the health profile of a pouch.",
        [("What flavour should I try first?", "Mint is a common starting point, but berry or citrus may suit shoppers who prefer fruit flavours."), ("Does flavour affect strength?", "Flavour can affect perceived intensity, but the mg strength should still be checked."), ("Where can I browse by flavour?", "Use Kangoo's flavour explorer and flavour landing pages.")],
    ),
    make_extra_guide(
        "Citrus Nicotine Pouches Guide: Lemon, Lime and Fresh Flavours",
        "citrus nicotine pouches",
        "Citrus Nicotine Pouches | Lemon, Lime and Fresh Picks",
        "Compare citrus nicotine pouches, including lemon, lime and fresh flavour profiles, with tips for choosing strength and brand.",
        ["lemon nicotine pouches", "lime nicotine pouches", "zyn citrus", "fruit nicotine pouches", "nicotine pouch flavours"],
        "Commercial investigation",
        [("flavour explorer", LINKS["flavour_explorer"]), ("ZYN Citrus", LINKS["zyn_citrus"]), ("shop nicotine pouches", LINKS["all"]), ("ZYN nicotine pouches", LINKS["zyn"])],
        "Citrus nicotine pouches are useful for shoppers who want a fresh, fruit-led flavour without choosing a heavy mint profile.",
        "Compare citrus options by exact flavour, strength, pouch format and brand availability.",
        "ZYN Citrus can be a useful product link when stocked, supported by the flavour explorer and full category.",
        "Keep recommendations tied to live stock and exact product strength.",
        [("Are citrus pouches usually strong?", "Not necessarily. Strength depends on the individual product."), ("Is citrus similar to mint?", "Both can feel fresh, but citrus is fruit-led rather than cooling-led."), ("Where should I compare citrus options?", "Use the flavour explorer and product-category pages.")],
    ),
    make_extra_guide(
        "Coffee Nicotine Pouches Guide: Who They Suit and What to Compare",
        "coffee nicotine pouches",
        "Coffee Nicotine Pouches | Flavour and Buying Guide",
        "A guide to coffee nicotine pouches, who may prefer them and how to compare flavour, strength and brand options at Kangoo.",
        ["coffee pouches", "coffee snus", "nicotine pouch flavours", "best nicotine pouches uk", "tobacco free nicotine pouches"],
        "Commercial investigation",
        [("flavour explorer", LINKS["flavour_explorer"]), ("shop nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("compare pouches", LINKS["compare"])],
        "Coffee nicotine pouches suit shoppers who want a warmer flavour profile than mint, ice or citrus.",
        "Compare coffee-style pouches by sweetness, strength, pouch size and whether the product is available in live stock.",
        "Use flavour explorer as the main route, then category links for current product options.",
        "Avoid implying caffeine content unless the product specifically contains caffeine and the product page confirms it.",
        [("Do coffee nicotine pouches contain caffeine?", "Not necessarily. Check the product page before assuming caffeine content."), ("Are coffee pouches tobacco-free?", "Many nicotine pouches are tobacco-free, but always check the product details."), ("Who might like coffee flavour?", "Adult users who prefer warmer, less icy flavour profiles may compare coffee-style pouches.")],
    ),
    make_extra_guide(
        "Ice Nicotine Pouches Guide: Cooling Mint and Freeze Flavours",
        "ice nicotine pouches",
        "Ice Nicotine Pouches | Cooling Mint and Freeze Guide",
        "Compare ice nicotine pouches, cooling mint flavours and freeze-style products by strength, pouch feel and brand.",
        ["ice nicotine pouches", "mint nicotine pouches", "velo freezing peppermint", "pablo ice cold", "strong nicotine pouches"],
        "Commercial investigation",
        [("mint nicotine pouches", LINKS["mint"]), ("VELO Freezing Peppermint", LINKS["velo_freezing"]), ("PABLO Ice Cold", LINKS["pablo_ice"]), ("strong nicotine pouches", LINKS["strong"])],
        "Ice nicotine pouches focus on cooling flavour, which can make a pouch feel sharper even before comparing mg strength.",
        "Compare ice pouches by cooling intensity, nicotine strength and brand, because mint and ice profiles can vary widely.",
        "Use mint and strong pages as supporting routes, with VELO and PABLO product links when live stock supports them.",
        "Do not assume a colder flavour means a higher nicotine strength.",
        [("Are ice pouches stronger?", "Not always. Cooling can affect perceived intensity, but strength depends on the product."), ("Which brands have ice-style options?", "VELO and PABLO can be useful comparison brands when stocked."), ("Should beginners choose ice pouches?", "Lower-strength mint options may be easier to compare first.")],
    ),
    make_extra_guide(
        "Mini Nicotine Pouches Guide: Smaller Format, Strength and Fit",
        "mini nicotine pouches",
        "Mini Nicotine Pouches | Format and Strength Guide",
        "Learn how mini nicotine pouches compare with slim or regular pouches, including format, comfort, flavour and strength checks.",
        ["zyn mini", "mini pouches", "nicotine pouch formats", "small nicotine pouches", "zyn pouches"],
        "Commercial investigation",
        [("ZYN nicotine pouches", LINKS["zyn"]), ("shop nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"])],
        "Mini nicotine pouches are smaller-format pouches that many shoppers compare for discretion and fit.",
        "Format does not automatically tell you strength, so check the exact mg and product details before buying.",
        "Use ZYN and full category pages as the main routes for mini-style searches.",
        "Avoid saying mini pouches are weaker or safer just because they are smaller.",
        [("Are mini pouches always low strength?", "No. Always check the exact strength on the product page."), ("Who might prefer mini pouches?", "Adult users who prefer a smaller pouch format may compare mini options."), ("Where can I find mini pouches?", "Start with the ZYN category and full nicotine pouch range.")],
    ),
    make_extra_guide(
        "Slim Nicotine Pouches Guide: Format, Feel and Buying Tips",
        "slim nicotine pouches",
        "Slim Nicotine Pouches | Format and Buying Tips",
        "Compare slim nicotine pouches by fit, flavour, strength and brand, with a simple Kangoo buying route.",
        ["slim snus", "slim pouches", "nicotine pouch formats", "strong nicotine pouches", "nicotine pouches uk"],
        "Commercial investigation",
        [("shop nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"]), ("compare pouches", LINKS["compare"])],
        "Slim nicotine pouches use a narrower format than some regular pouches, which can change how they feel under the lip.",
        "Format should be compared alongside strength, moisture and flavour because pouch shape is only one buying factor.",
        "Send shoppers into the full category and pouch finder to compare format and strength together.",
        "Avoid turning format into a health or safety claim.",
        [("Are slim pouches more discreet?", "They may feel more discreet for some users, but comfort is personal."), ("Are slim pouches strong?", "Some are, some are not. Check exact strength."), ("How do I compare pouch formats?", "Use product details and the pouch finder to compare format, flavour and strength.")],
    ),
    make_extra_guide(
        "3mg Nicotine Pouches Guide: Lower-Strength Options Explained",
        "3mg nicotine pouches",
        "3mg Nicotine Pouches | Lower-Strength Guide",
        "Compare 3mg nicotine pouches, who lower-strength options may suit and how to browse mint, berry and ZYN-style picks.",
        ["3mg pouches", "low strength nicotine pouches", "zyn 3mg", "mini nicotine pouches", "nicotine pouch strength"],
        "Commercial investigation",
        [("strength ladder", LINKS["strength_ladder"]), ("ZYN Cool Mint", LINKS["zyn_cool"]), ("ZYN Black Cherry", LINKS["zyn_cherry"]), ("shop nicotine pouches", LINKS["all"])],
        "3mg nicotine pouches are typically treated as lower-strength options compared with strong and extra-strong products.",
        "They are useful for shoppers who want to compare flavour and brand without jumping straight to higher-strength pouches.",
        "Use the strength ladder as the hub, then link to ZYN-style lower-strength products when stocked.",
        "Do not describe lower strength as safe; it still contains addictive nicotine.",
        [("Is 3mg low strength?", "It is generally lower than strong and extra-strong pouch options."), ("Is 3mg only for beginners?", "Not necessarily. Some experienced users prefer lower-strength pouches."), ("Where should I compare 3mg pouches?", "Use the strength ladder and product category filters.")],
    ),
    make_extra_guide(
        "6mg Nicotine Pouches Guide: Balanced Strength Buying Tips",
        "6mg nicotine pouches",
        "6mg Nicotine Pouches | Balanced Strength Guide",
        "Compare 6mg nicotine pouches as a balanced-strength option, with buying tips for mint, citrus, berry and brand choice.",
        ["6mg pouches", "zyn 6mg", "velo 6mg", "nicotine pouch strength", "balanced nicotine pouches"],
        "Commercial investigation",
        [("strength ladder", LINKS["strength_ladder"]), ("VELO Bright Spearmint", LINKS["velo_bright"]), ("ZYN Citrus", LINKS["zyn_citrus"]), ("shop nicotine pouches", LINKS["all"])],
        "6mg nicotine pouches often sit in a balanced range for adult users comparing everyday pouch options.",
        "Strength feel can still vary by brand, flavour and format, so 6mg should not be treated as identical across products.",
        "Use the strength ladder, then link to relevant ZYN or VELO products when stocked.",
        "Keep the article clear that nicotine remains addictive at every strength.",
        [("Is 6mg strong?", "It is often a balanced or mid-level option, but product feel varies."), ("Which flavours come in 6mg?", "Mint, citrus and other flavours may be available depending on live stock."), ("Should I choose 6mg or 3mg?", "Compare your experience level and preferred pouch feel before choosing.")],
    ),
    make_extra_guide(
        "11mg Nicotine Pouches Guide: Stronger Pouch Options",
        "11mg nicotine pouches",
        "11mg Nicotine Pouches | Stronger Option Guide",
        "Compare 11mg nicotine pouches, stronger mint options and how to check whether a higher-strength pouch suits your preference.",
        ["11mg pouches", "velo 11mg", "strong nicotine pouches", "velo freezing peppermint", "nicotine pouch strength"],
        "Commercial investigation",
        [("strong nicotine pouches", LINKS["strong"]), ("VELO Freezing Peppermint", LINKS["velo_freezing"]), ("strength ladder", LINKS["strength_ladder"]), ("shop nicotine pouches", LINKS["all"])],
        "11mg nicotine pouches are usually in stronger territory and should be compared carefully by experienced adult users.",
        "Cooling flavours can make stronger pouches feel sharper, so compare mg strength and flavour profile together.",
        "Route readers through the strong page and strength ladder, with VELO Freezing Peppermint as a relevant stocked example when available.",
        "Avoid suggesting higher strength is automatically better value.",
        [("Is 11mg suitable for beginners?", "Usually no. It is better suited to experienced adult nicotine users."), ("Does mint make 11mg feel stronger?", "Cooling can affect perceived intensity, but the mg still matters."), ("Where can I compare 11mg pouches?", "Use strong nicotine pouch pages and the strength ladder.")],
    ),
    make_extra_guide(
        "30mg Nicotine Pouches Guide: Extra Strong Options and Cautions",
        "30mg nicotine pouches",
        "30mg Nicotine Pouches | Extra Strong Guide",
        "A cautious guide to 30mg nicotine pouches, extra strong products and why high-strength pouches suit experienced adult users only.",
        ["extra strong nicotine pouches", "pablo grape ice", "strongest nicotine pouches", "30mg pouches", "pablo nicotine pouches"],
        "Commercial investigation",
        [("extra strong nicotine pouches", LINKS["extra"]), ("PABLO Grape Ice", LINKS["pablo_grape"]), ("PABLO nicotine pouches", LINKS["pablo"]), ("strength ladder", LINKS["strength_ladder"])],
        "30mg nicotine pouches sit at the extra-strong end of the category and should be approached only by experienced adult nicotine users.",
        "The practical comparison is not just 'highest mg' but whether the strength, flavour and pouch format suit the user.",
        "Use extra-strong and PABLO pages as the main route, with specific product links only where stock is live.",
        "Keep the caution prominent and avoid presenting 30mg products as a general recommendation.",
        [("Are 30mg pouches for beginners?", "No. They are extra-strong products for experienced adult nicotine users."), ("Is 30mg always better value?", "No. A pouch that is too strong is not a better buy."), ("Where can I compare extra strong products?", "Use Kangoo's extra strong category and strength ladder.")],
    ),
    make_extra_guide(
        "Cheapest Snus UK? What Shoppers Actually Mean",
        "cheapest snus uk",
        "Cheapest Snus UK? Tobacco-Free Pouch Guide",
        "Many UK shoppers search for cheapest snus. Learn how to compare tobacco-free nicotine pouches, 99p trials and pack pricing instead.",
        ["cheap snus uk", "buy snus uk", "cheap nicotine pouches uk", "99p nicotine pouches", "snus uk"],
        "Commercial investigation",
        [("99p nicotine pouches", LINKS["trial"]), ("shop nicotine pouches", LINKS["all"]), ("compare pouches", LINKS["compare"]), ("ZYN nicotine pouches", LINKS["zyn"])],
        "When UK shoppers search for cheapest snus, they are often looking for low-price tobacco-free nicotine pouches.",
        "The best comparison is trial price, regular can price, pack pricing, live stock and delivery threshold.",
        "Use 99p pouches for the price hook, then send repeat buyers to the full nicotine pouch range.",
        "Avoid absolute cheapest claims unless the evidence is current, dated and verifiable.",
        [("Can Kangoo say cheapest snus?", "Only with current evidence. Safer wording is from 99p trial pouches or cheap nicotine pouch options."), ("Is cheap snus the same as cheap nicotine pouches?", "Often shoppers mean nicotine pouches, but traditional snus is different."), ("Where should price-focused users start?", "Start with 99p pouches, then compare regular cans and pack pricing.")],
    ),
    make_extra_guide(
        "Nicotine Pouches Online UK: Delivery, Price and Buying Checklist",
        "nicotine pouches online uk",
        "Nicotine Pouches Online UK | Delivery and Buying Guide",
        "Shop nicotine pouches online in the UK with a practical checklist for delivery, discreet packaging, price, strength and brand choice.",
        ["buy nicotine pouches online", "nicotine pouches uk", "online snus uk", "cheap nicotine pouches", "discreet delivery nicotine pouches"],
        "Transactional",
        [("shop nicotine pouches", LINKS["all"]), ("99p nicotine pouches", LINKS["trial"]), ("compare pouches", LINKS["compare"]), ("pouch finder", LINKS["finder"])],
        "Buying nicotine pouches online is easiest when shoppers compare live stock, price, strength, flavour and delivery details in one place.",
        "Online buyers should check product strength, pack pricing, delivery threshold and whether the product is a trial or regular can.",
        "Route users to the full category, 99p pouch range and finder depending on their intent.",
        "Keep delivery claims accurate and aligned with the WooCommerce settings.",
        [("Can I buy nicotine pouches online in the UK?", "Adult shoppers can buy tobacco-free nicotine pouches online from retailers such as Kangoo."), ("What should I check before ordering?", "Check strength, flavour, stock, delivery terms and total basket value."), ("Are online prices fixed?", "No. Pricing and availability can change, so check the live product page.")],
    ),
    make_extra_guide(
        "Nicotine Pouches Near Me vs Online: Which Buying Route Works Best?",
        "nicotine pouches near me",
        "Nicotine Pouches Near Me vs Online | Buying Guide",
        "Compare searching for nicotine pouches near me with buying online, including stock range, price checks, delivery and discreet packaging.",
        ["where to buy nicotine pouches", "tobacco shop near me", "vape shops near me", "nicotine pouches uk", "buy snus online"],
        "Commercial investigation",
        [("shop nicotine pouches", LINKS["all"]), ("99p nicotine pouches", LINKS["trial"]), ("compare pouches", LINKS["compare"]), ("ZYN nicotine pouches", LINKS["zyn"])],
        "Near me searches are usually about convenience, but online shopping can offer a clearer way to compare range, strength and price.",
        "A local shop may be useful for immediate need, while online categories make it easier to compare brands, flavours and pack pricing.",
        "Use the full category as the main conversion route and 99p pouches as the value hook.",
        "Avoid claims that every local shop is more expensive; frame this as range and comparison convenience.",
        [("Is online better than local shops?", "It depends on urgency, stock and price. Online is usually easier for comparing range."), ("Can I compare brands online?", "Yes. Category pages make it easier to compare ZYN, VELO, PABLO and KILLA."), ("What if I need pouches today?", "A local shop may be faster, but online buying gives clearer product comparison.")],
    ),
    make_extra_guide(
        "Tobacco-Free Snus UK: What the Phrase Means",
        "tobacco free snus uk",
        "Tobacco-Free Snus UK | Nicotine Pouch Guide",
        "Learn what UK shoppers mean by tobacco-free snus and how it relates to tobacco-free nicotine pouches at Kangoo.",
        ["tobacco free snus", "tobacco free nicotine pouches", "snus uk", "nicotine pouches uk", "oral nicotine products"],
        "Informational and commercial",
        [("shop nicotine pouches", LINKS["all"]), ("ZYN nicotine pouches", LINKS["zyn"]), ("VELO nicotine pouches", LINKS["velo"]), ("pouch finder", LINKS["finder"])],
        "Tobacco-free snus is a phrase many shoppers use when they mean tobacco-free nicotine pouches.",
        "The important distinction is that traditional snus contains tobacco, while nicotine pouches generally do not contain tobacco leaf.",
        "Route this topic to tobacco-free nicotine pouch categories and explain the terminology clearly.",
        "Avoid mixing legal claims about traditional snus with claims about tobacco-free pouches.",
        [("Is tobacco-free snus the same as nicotine pouches?", "In many searches, yes, but the clearer product term is tobacco-free nicotine pouches."), ("Does tobacco-free mean nicotine-free?", "No. Tobacco-free nicotine pouches still contain nicotine unless marked otherwise."), ("Where can I compare tobacco-free pouches?", "Use Kangoo's nicotine pouch category and finder tools.")],
    ),
    make_extra_guide(
        "Oral Nicotine Products: Pouches, Gum, Lozenges and Vapes Compared",
        "oral nicotine products",
        "Oral Nicotine Products | Pouches, Gum and Lozenges",
        "Compare oral nicotine products such as pouches, gum and lozenges, plus how they differ from vapes and tobacco products.",
        ["oral nicotine", "oral nicotine pouches", "smokeless nicotine", "nicotine replacement therapy", "nicotine pouches uk"],
        "Informational",
        [("shop nicotine pouches", LINKS["all"]), ("what are nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("compare pouches", LINKS["compare"])],
        "Oral nicotine products can include several different formats, but they do not all have the same status, use method or purpose.",
        "A nicotine pouch is a retail pouch product, while gum and lozenges may be licensed NRT products depending on the product.",
        "Use this guide as an educational bridge to the nicotine pouch category, not as medical advice.",
        "Keep distinctions clear and avoid claiming pouches are licensed stop-smoking products.",
        [("Are all oral nicotine products the same?", "No. Format, regulation and intended use differ."), ("Are pouches smoked or vaped?", "No. They are placed under the lip."), ("Are nicotine pouches NRT?", "No, they are not licensed stop-smoking products in the UK.")],
    ),
    make_extra_guide(
        "Smokeless Nicotine: What It Means and How Pouches Compare",
        "smokeless nicotine",
        "Smokeless Nicotine | How Nicotine Pouches Compare",
        "Understand smokeless nicotine searches, how nicotine pouches fit the category and how adult users can compare pouch options.",
        ["smokeless nicotine pouches", "smokeless tobacco", "nicotine pouches uk", "tobacco free nicotine", "oral nicotine products"],
        "Informational",
        [("shop nicotine pouches", LINKS["all"]), ("tobacco-free nicotine pouches", LINKS["all"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"])],
        "Smokeless nicotine is a broad search phrase, and nicotine pouches are one tobacco-free product format people may be looking for.",
        "Pouches should be compared separately from smokeless tobacco because the product composition and legal context can differ.",
        "Route users to tobacco-free nicotine pouch education, then category and finder pages.",
        "Avoid implying smokeless means harmless.",
        [("Does smokeless mean risk-free?", "No. Nicotine is addictive and smokeless should not be presented as safe."), ("Are nicotine pouches tobacco-free?", "Most modern pouches sold by Kangoo are tobacco-free, but check product details."), ("How should I compare smokeless nicotine options?", "Compare format, strength, intended use and product information.")],
    ),
    make_extra_guide(
        "Nicotine Pouch Brands UK: ZYN, VELO, PABLO and KILLA Compared",
        "nicotine pouch brands uk",
        "Nicotine Pouch Brands UK | ZYN, VELO, PABLO, KILLA",
        "Compare nicotine pouch brands in the UK, including ZYN, VELO, PABLO and KILLA by flavour, strength and buying route.",
        ["best nicotine pouch brands", "zyn vs velo", "pablo nicotine pouches", "killa nicotine pouches", "velo nicotine pouches"],
        "Commercial investigation",
        [("ZYN nicotine pouches", LINKS["zyn"]), ("VELO nicotine pouches", LINKS["velo"]), ("PABLO nicotine pouches", LINKS["pablo"]), ("KILLA nicotine pouches", LINKS["killa"]), ("compare pouches", LINKS["compare"])],
        "The best nicotine pouch brand depends on flavour preference, strength range and whether the shopper wants everyday or stronger products.",
        "ZYN and VELO are useful mainstream comparisons, while PABLO and KILLA are more relevant to stronger-product shoppers.",
        "Use brand categories as the main internal links and compare tools for users who need help deciding.",
        "Avoid ranking brands without a dated method or evidence.",
        [("Which brand should I try first?", "ZYN and VELO are useful starting comparisons for many shoppers."), ("Which brands are strongest?", "PABLO and KILLA often appear in stronger comparisons, but check each product page."), ("Can I compare by flavour?", "Yes. Use the flavour explorer or brand pages.")],
    ),
    make_extra_guide(
        "Best ZYN Flavours UK: Mint, Citrus and Berry Picks",
        "best zyn flavours",
        "Best ZYN Flavours UK | Mint, Citrus and Berry Guide",
        "Compare ZYN flavours including Cool Mint, Citrus and Black Cherry, with strength and buying tips for Kangoo shoppers.",
        ["zyn flavours", "zyn cool mint", "zyn black cherry", "zyn citrus", "zyn pouches"],
        "Commercial investigation",
        [("ZYN nicotine pouches", LINKS["zyn"]), ("ZYN Cool Mint", LINKS["zyn_cool"]), ("ZYN Black Cherry", LINKS["zyn_cherry"]), ("ZYN Citrus", LINKS["zyn_citrus"]), ("flavour explorer", LINKS["flavour_explorer"])],
        "The best ZYN flavour depends on whether the shopper wants mint freshness, fruit sweetness or citrus sharpness.",
        "Compare flavour and strength together, because the same brand can include multiple pouch styles.",
        "Use the ZYN category as the hub and product links for stocked flavour pages.",
        "Do not overstate a best flavour as universal; make it preference-led.",
        [("What is the best ZYN flavour?", "It depends on taste. Cool Mint, Black Cherry and Citrus are useful comparison points."), ("Does flavour change strength?", "The mg strength is separate, but flavour can affect perceived intensity."), ("Where can I browse ZYN flavours?", "Use Kangoo's ZYN category and flavour explorer.")],
    ),
    make_extra_guide(
        "Best VELO Flavours UK: Spearmint, Peppermint and Berry",
        "best velo flavours",
        "Best VELO Flavours UK | Mint and Berry Guide",
        "Compare VELO flavours including Bright Spearmint, Freezing Peppermint and Ruby Berry, with strength and buying tips.",
        ["velo flavours", "velo freezing peppermint", "velo ruby berry", "velo bright spearmint", "velo pouches"],
        "Commercial investigation",
        [("VELO nicotine pouches", LINKS["velo"]), ("VELO Bright Spearmint", LINKS["velo_bright"]), ("VELO Freezing Peppermint", LINKS["velo_freezing"]), ("VELO Ruby Berry", LINKS["velo_ruby"]), ("flavour explorer", LINKS["flavour_explorer"])],
        "VELO flavour comparisons usually split between fresh mint, stronger cooling and fruit-led berry profiles.",
        "Shoppers should compare flavour family first, then check exact strength and pouch format.",
        "Use VELO as the brand hub, supported by product links when live stock is available.",
        "Avoid saying a flavour is objectively best for everyone.",
        [("Which VELO flavour should I try?", "Bright Spearmint and Freezing Peppermint suit mint shoppers, while Ruby Berry suits fruit-led preferences."), ("Are VELO flavours all the same strength?", "No. Check the exact product page."), ("Where can I compare VELO flavours?", "Use the VELO category and flavour explorer.")],
    ),
    make_extra_guide(
        "Pouch Finder Guide: How to Choose Strength and Flavour at Kangoo",
        "pouch finder",
        "Pouch Finder Guide | Choose Strength and Flavour",
        "Learn how to use the Kangoo Pouch Finder to choose nicotine pouch strength, flavour, brand and format without guessing.",
        ["find nicotine pouches", "best nicotine pouches uk", "nicotine pouch strength", "nicotine pouch flavours", "compare nicotine pouches"],
        "Commercial investigation",
        [("Kangoo Pouch Finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"]), ("flavour explorer", LINKS["flavour_explorer"]), ("compare pouches", LINKS["compare"]), ("shop nicotine pouches", LINKS["all"])],
        "A pouch finder helps shoppers move from broad research into a smaller set of products based on strength and flavour preference.",
        "This guide should explain when to use finder, compare, strength ladder and flavour explorer tools.",
        "Use tool pages as the main internal links, then direct confident shoppers to the full product category.",
        "Avoid pretending the finder gives medical advice; it is a shopping helper.",
        [("What does the pouch finder do?", "It helps shoppers narrow by strength, flavour and experience level."), ("Is the pouch finder medical advice?", "No. It is a shopping tool."), ("What if I already know my brand?", "Go directly to the relevant brand page such as ZYN or VELO.")],
    ),
    make_extra_guide(
        "Discreet Nicotine Pouches: Online Buying and Delivery Guide",
        "discreet nicotine pouches",
        "Discreet Nicotine Pouches | Online Delivery Guide",
        "Compare discreet nicotine pouch buying online, including delivery, packaging, strength checks and product choice at Kangoo.",
        ["discreet delivery nicotine pouches", "buy nicotine pouches online", "nicotine pouches online uk", "tobacco free pouches", "nicotine pouch"],
        "Transactional",
        [("shop nicotine pouches", LINKS["all"]), ("99p nicotine pouches", LINKS["trial"]), ("pouch finder", LINKS["finder"]), ("compare pouches", LINKS["compare"])],
        "Discreet buying intent is usually about privacy, packaging, delivery and a simple product choice.",
        "The page should connect discreet delivery copy with real buying decisions: strength, flavour, price and stock.",
        "Use commercial category links early because this is close to purchase intent.",
        "Keep delivery wording aligned with actual fulfilment terms and the live free-delivery threshold.",
        [("Does Kangoo offer discreet packaging?", "Kangoo product and delivery copy should state current packaging terms clearly on site."), ("What should I check before ordering?", "Check strength, flavour, stock and delivery threshold."), ("Can I try before buying more?", "Use the 99p pouch range when trial stock is available.")],
    ),
    make_extra_guide(
        "Nicotine Pouches for Beginners: First Order Checklist",
        "nicotine pouches for beginners",
        "Nicotine Pouches for Beginners | First Order Checklist",
        "A beginner-friendly checklist for adult nicotine users comparing nicotine pouch strength, flavour, brand and first-order value.",
        ["beginner nicotine pouches", "what are nicotine pouches", "how to use nicotine pouches", "low strength nicotine pouches", "99p nicotine pouches"],
        "Informational and commercial",
        [("99p nicotine pouches", LINKS["trial"]), ("pouch finder", LINKS["finder"]), ("strength ladder", LINKS["strength_ladder"]), ("flavour explorer", LINKS["flavour_explorer"]), ("shop nicotine pouches", LINKS["all"])],
        "Beginners should make a small, careful first order based on lower or balanced strength, familiar flavour and live stock.",
        "The first decision should be strength comfort, then flavour, then brand and price.",
        "Start with 99p trials and finder tools before recommending larger pack pricing.",
        "Keep the page for adult nicotine users and do not encourage non-users to start.",
        [("Should beginners choose extra strong pouches?", "No. Extra strong products are better suited to experienced adult nicotine users."), ("What flavour should beginners try?", "Mint or berry are common starting points, depending on taste."), ("Is a 99p pouch useful for beginners?", "Yes, when available, it can help sample a flavour or strength before a regular order.")],
    ),
    make_extra_guide(
        "Are Nicotine Pouches Tobacco-Free? UK Product Guide",
        "are nicotine pouches tobacco free",
        "Are Nicotine Pouches Tobacco-Free? UK Guide",
        "Learn whether nicotine pouches are tobacco-free, how they differ from traditional snus and what to check on Kangoo product pages.",
        ["tobacco free nicotine pouches", "tobacco free snus", "what are nicotine pouches", "snus vs nicotine pouches", "nicotine pouches uk"],
        "Informational",
        [("tobacco-free nicotine pouches", LINKS["all"]), ("ZYN nicotine pouches", LINKS["zyn"]), ("VELO nicotine pouches", LINKS["velo"]), ("shop nicotine pouches", LINKS["all"])],
        "Modern nicotine pouches are generally tobacco-free, but they still contain nicotine unless stated otherwise.",
        "The clearest comparison is tobacco-free nicotine pouch versus traditional tobacco snus.",
        "Use this guide to support tobacco-free category messaging and link to brand pages.",
        "Do not imply tobacco-free means nicotine-free or risk-free.",
        [("Do nicotine pouches contain tobacco?", "Most modern pouches sold by Kangoo are tobacco-free, but check individual product details."), ("Do tobacco-free pouches contain nicotine?", "Yes, unless the product says otherwise."), ("Are tobacco-free pouches the same as snus?", "No. Traditional snus contains tobacco.")],
    ),
]

for offset, guide in enumerate(EXTRA_GUIDES):
    guide["scheduled_at"] = (EXTRA_PUBLISH_START + timedelta(days=offset)).strftime("%Y-%m-%d 09:00:00")

GUIDES.extend(EXTRA_GUIDES)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F5FBF")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_document_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 21, "111111"),
        ("Subtitle", 11, "666666"),
        ("Heading 1", 16, "111111"),
        ("Heading 2", 13, "111111"),
        ("Heading 3", 11, "111111"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_header_footer(doc: Document, label: str):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    header = section.header
    p = header.paragraphs[0]
    p.text = label
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Kangoo Pouches SEO guide"
    fp.runs[0].font.name = "Arial"
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_link_list(doc: Document, items: list[tuple[str, str]]):
    for label, url in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ")
        add_hyperlink(p, url, url)


def write_guide(index: int, guide: dict) -> Path:
    doc = Document()
    set_document_styles(doc)
    add_header_footer(doc, guide["focus"])

    doc.add_paragraph(guide["title"], style="Title")
    doc.add_paragraph("Ready-to-paste WordPress blog draft and SEO brief", style="Subtitle")

    doc.add_heading("SEO Brief", level=1)
    brief_items = [
        f"Suggested slug: /blog/{slugify(guide['title'])}/",
        f"Focus keyphrase: {guide['focus']}",
        f"Search intent: {guide['intent']}",
        f"Keyword evidence: {keyword_note(guide['focus'])}",
        "Focus keyphrase usage: 4-6 natural mentions across H1/title, opening paragraph, one H2 if natural, body and meta. Do not force exact-match repetition.",
        "Secondary keyword usage: 1-2 natural mentions each, mainly in subheadings, body copy or FAQs.",
        "Compliance note: for adults who already use nicotine products; avoid medical, stop-smoking or absolute lowest-price claims.",
    ]
    add_bullets(doc, brief_items)

    doc.add_heading("Meta Title", level=2)
    doc.add_paragraph(guide["meta_title"])
    doc.add_heading("Meta Description", level=2)
    doc.add_paragraph(guide["meta_description"])
    doc.add_heading("Secondary Keywords", level=2)
    add_bullets(doc, [f"{kw} - {keyword_note(kw)}" for kw in guide["secondary"]])
    doc.add_heading("Internal Link Plan", level=2)
    add_link_list(doc, guide["links"])

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("WordPress Post Draft", level=1)
    doc.add_paragraph(guide["title"], style="Title")
    doc.add_paragraph(NICOTINE_NOTE)

    intro = (
        f"If you are searching for {guide['focus']}, the useful answer is not just a product name. "
        f"You need a clear way to compare price, strength, flavour, stock and brand fit before you add a pouch to your basket. "
        f"This guide explains the practical route for Kangoo shoppers and points you toward the most relevant category, tool and product pages."
    )
    doc.add_paragraph(intro)

    for heading, paragraphs in guide["sections"]:
        doc.add_heading(heading, level=2)
        for paragraph in paragraphs:
            doc.add_paragraph(public_text(paragraph))

    doc.add_heading("How to choose between similar options", level=2)
    doc.add_paragraph(
        "When two pouches look similar, start with the product details rather than the product image. "
        "Check the exact strength, pouch format, flavour family, price and stock status before deciding."
    )
    doc.add_paragraph(
        "For flavour-led choices, compare mint, berry, citrus, coffee, tropical and ice profiles first. "
        "For strength-led choices, use the strength ladder so you can compare light, balanced, strong and extra-strong products more clearly."
    )

    doc.add_heading("Price, stock and delivery checks", level=2)
    doc.add_paragraph(
        "A good value order is not only about the lowest single pouch price. Trial pouches can help with sampling, while regular cans and selected pack pricing are better suited to repeat orders."
    )
    doc.add_paragraph(
        "Before ordering, check that the product is in stock, the strength matches your preference, and delivery terms still match the current Kangoo checkout."
    )

    doc.add_heading("Buying checklist", level=2)
    add_bullets(doc, [
        "Check the exact strength on the product page before buying.",
        "Compare flavour family first: mint, berry, citrus, coffee, tropical or ice.",
        "Check whether the product is in stock before placing an order.",
        "Use category pages for evergreen browsing and product pages for live picks.",
        "Keep wording factual: price, stock, strength, flavour, pack size and delivery terms.",
    ])

    doc.add_heading("Recommended Kangoo links to add in the body", level=2)
    add_link_list(doc, guide["links"])

    doc.add_heading("FAQs", level=2)
    for question, answer in guide["faq"]:
        doc.add_heading(question, level=3)
        doc.add_paragraph(public_text(answer))

    doc.add_heading("Editor Notes Before Publishing", level=1)
    add_bullets(doc, [
        "Check live stock before adding individual product cards or buttons.",
        "Do not add unsupported comparison claims such as 'cheapest in the UK' unless you keep verifiable evidence and dates.",
        "Add the focus keyphrase to the WordPress SEO title, slug, first paragraph and one subheading only if it reads naturally.",
        "Use Product schema only on product pages; this blog should use article/blog schema from WordPress or Yoast.",
        "Add one featured image if available, but keep the first paragraph crawlable text.",
    ])

    doc.add_heading("External Reference Notes", level=2)
    reference_items = [
        ("NHS: nicotine pouches are not licensed stop-smoking products in the UK", SOURCES["nhs"]),
        ("GOV.UK: oral snuff/snus sale is prohibited in the UK", SOURCES["gov_snus"]),
        ("GOV.UK: marketing must be truthful and accurate", SOURCES["gov_ads"]),
        ("ASA/CAP: comparison claims should be verifiable", SOURCES["asa_compare"]),
    ]
    if "snus" not in guide["focus"].lower() and "legal" not in guide["focus"].lower():
        reference_items = [reference_items[0], reference_items[2], reference_items[3]]
    add_link_list(doc, reference_items)

    filename = f"{index:02d}-{slugify(guide['title'])}.docx"
    path = OUT_DIR / filename
    doc.save(path)
    return path


def write_plan(paths: list[Path]) -> Path:
    doc = Document()
    set_document_styles(doc)
    add_header_footer(doc, "Kangoo blog publishing plan")
    doc.add_paragraph("Kangoo Blog Publishing Plan", style="Title")
    doc.add_paragraph("SEO briefs, release order and internal-link plan for the generated guide batch.", style="Subtitle")

    doc.add_heading("How to publish", level=1)
    add_numbered(doc, [
        "Publish one guide per day for the first 18 days, ideally at the same time each morning.",
        "Start with commercial pages first: cheap, 99p, best, strength and strongest pouches.",
        "Then publish education pages: how to use, what is snus, snus vs nicotine pouches and legal/tobacco-free explainers.",
        "Finish with brand and flavour pages, then update older posts to link to the newly published guides.",
        "After each post is live, request indexing for the post and make sure it links back to one commercial category page.",
    ])

    doc.add_heading("Release schedule", level=1)
    for i, guide in enumerate(GUIDES, start=1):
        scheduled_at = guide.get("scheduled_at")
        publish_date = scheduled_at[:10] if scheduled_at else (INITIAL_PUBLISH_START + timedelta(days=i - 1)).isoformat()
        doc.add_heading(f"Day {i}: {guide['title']}", level=2)
        add_bullets(doc, [
            f"Suggested publish date: {publish_date}",
            f"Focus keyphrase: {guide['focus']}",
            f"Meta title: {guide['meta_title']}",
            f"Meta description: {guide['meta_description']}",
            f"Primary intent: {guide['intent']}",
            f"Guide file: {paths[i - 1].name}",
        ])
        doc.add_paragraph("Internal links:")
        add_link_list(doc, guide["links"][:5])

    doc.add_heading("Priority groups", level=1)
    add_bullets(doc, [
        "Highest commercial priority: Cheap Nicotine Pouches UK, 99p Nicotine Pouches UK, Best Nicotine Pouches UK.",
        "Highest topical authority priority: Nicotine Pouch Strength Guide, What Is Snus, Snus vs Nicotine Pouches, How To Use Nicotine Pouches.",
        "Brand priority: ZYN, VELO, ZYN vs VELO, PABLO, KILLA.",
        "Long-tail flavour priority: Mint Nicotine Pouches, Berry Nicotine Pouches.",
        "Trust/support priority: Legal UK guide, Tobacco-Free Nicotine Pouches, Nicotine Pouches vs Vapes.",
    ])

    doc.add_heading("Internal link rules", level=1)
    add_bullets(doc, [
        "Every post should link to exactly one main money page in the first third of the article.",
        "Use category pages for evergreen links and product links only when stock is reliable.",
        "Add links both ways: once a post is live, add a small link from relevant category SEO copy or future blogs back to that guide.",
        "Avoid too many product links in educational posts. Two to four internal links is usually enough; commercial guides can use five to eight.",
        "Use natural anchors such as '99p nicotine pouches', 'ZYN nicotine pouches' and 'strength ladder'.",
    ])

    path = OUT_DIR / "00-kangoo-blog-publishing-plan.docx"
    doc.save(path)
    return path


def guide_topic(guide: dict) -> str:
    title = guide["title"].lower()

    if any(word in title for word in ("zyn", "velo", "pablo", "killa")):
        return "Brand Guide"

    if any(word in title for word in ("mint", "berry")):
        return "Flavour Guide"

    if any(word in title for word in ("strength", "strongest", "strong ", "extra strong")):
        return "Strength Guide"

    if any(word in title for word in ("snus", "legal", "tobacco-free", "vapes", "how to use")):
        return "Education"

    return "Buying Guide"


def content_html(guide: dict) -> str:
    parts = []
    parts.append(f"<p>{html.escape(NICOTINE_NOTE)}</p>")
    parts.append(
        "<p>"
        + html.escape(
            f"If you are searching for {guide['focus']}, the useful answer is not just a product name. "
            "You need a clear way to compare price, strength, flavour, stock and brand fit before you add a pouch to your basket. "
            "This guide explains the practical route for Kangoo shoppers and points you toward the most relevant category, tool and product pages."
        )
        + "</p>"
    )

    for heading, paragraphs in guide["sections"]:
        parts.append(f"<h2>{html.escape(heading)}</h2>")
        for paragraph in paragraphs:
            parts.append(f"<p>{html.escape(public_text(paragraph))}</p>")

    parts.append("<h2>How to choose between similar options</h2>")
    for paragraph in [
        "When two pouches look similar, start with the product details rather than the product image. Check the exact strength, pouch format, flavour family, price and stock status before deciding.",
        "For flavour-led choices, compare mint, berry, citrus, coffee, tropical and ice profiles first. For strength-led choices, use the strength ladder so you can compare light, balanced, strong and extra-strong products more clearly.",
    ]:
        parts.append(f"<p>{html.escape(paragraph)}</p>")

    parts.append("<h2>Price, stock and delivery checks</h2>")
    for paragraph in [
        "A good value order is not only about the lowest single pouch price. Trial pouches can help with sampling, while regular cans and selected pack pricing are better suited to repeat orders.",
        "Before ordering, check that the product is in stock, the strength matches your preference, and delivery terms still match the current Kangoo checkout.",
    ]:
        parts.append(f"<p>{html.escape(paragraph)}</p>")

    parts.append("<h2>Buying checklist</h2>")
    parts.append("<ul>")
    for item in [
        "Check the exact strength on the product page before buying.",
        "Compare flavour family first: mint, berry, citrus, coffee, tropical or ice.",
        "Check whether the product is in stock before placing an order.",
        "Use category pages for evergreen browsing and product pages for live picks.",
        "Keep wording factual: price, stock, strength, flavour, pack size and delivery terms.",
    ]:
        parts.append(f"<li>{html.escape(item)}</li>")
    parts.append("</ul>")

    parts.append("<h2>Recommended Kangoo links</h2>")
    parts.append("<ul>")
    for label, url in guide["links"]:
        parts.append(f'<li><a href="{html.escape(url, quote=True)}">{html.escape(label)}</a></li>')
    parts.append("</ul>")

    parts.append("<h2>Frequently Asked Questions</h2>")
    for question, answer in guide["faq"]:
        parts.append(f"<h3>{html.escape(question)}</h3>")
        parts.append(f"<p>{html.escape(public_text(answer))}</p>")

    return "\n".join(parts)


def write_seed_json() -> Path:
    seed_items = []
    html_dir = OUT_DIR / "html"
    html_dir.mkdir(parents=True, exist_ok=True)

    for index, guide in enumerate(GUIDES, start=1):
        body = content_html(guide)
        words = len(re.sub(r"<[^>]+>", " ", body).split())
        seed_item = {
            "order": index,
            "title": guide["title"],
            "slug": slugify(guide["title"]),
            "focus_keyphrase": guide["focus"],
            "seo_title": guide["meta_title"],
            "meta_description": guide["meta_description"],
            "standfirst": guide["meta_description"],
            "eyebrow": guide_topic(guide),
            "read_time": max(4, round(words / 180)),
            "topic": guide_topic(guide),
            "secondary_keywords": guide["secondary"],
            "content_html": body,
        }

        if guide.get("scheduled_at"):
            seed_item["scheduled_at"] = guide["scheduled_at"]

        seed_items.append(seed_item)

        html_path = html_dir / f"{index:02d}-{slugify(guide['title'])}.html"
        html_path.write_text(
            "\n".join(
                [
                    f"<!-- Title: {guide['title']} -->",
                    f"<!-- Focus keyphrase: {guide['focus']} -->",
                    f"<!-- SEO title: {guide['meta_title']} -->",
                    f"<!-- Meta description: {guide['meta_description']} -->",
                    f"<!-- Scheduled at: {guide.get('scheduled_at', 'Use seeder start date')} -->",
                    body,
                    "",
                ]
            ),
            encoding="utf-8",
        )

    path = OUT_DIR / "kangoo-blog-seed-data.json"
    path.write_text(json.dumps(seed_items, indent=2, ensure_ascii=False), encoding="utf-8")

    return path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = []
    for index, guide in enumerate(GUIDES, start=1):
        paths.append(write_guide(index, guide))
    plan = write_plan(paths)
    seed_json = write_seed_json()
    print(plan)
    for path in paths:
        print(path)
    print(seed_json)


if __name__ == "__main__":
    main()
